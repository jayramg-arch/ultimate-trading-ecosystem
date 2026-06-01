"""
Recovery Trading Manual — Generator Part 5
BULL MARKET SECTION: Sections 18–21
- Section 18: Philosophy & Unified Method
- Section 19: The Six Bull Market Catalysts
- Section 20: Strategy v4.5 — Complete Reference
- Section 21: Ultimate Screener v3.6 — Complete Reference
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = r"C:\Users\jayra\Documents\GeminiVSCode\Recovery_Trading_Manual.docx"
doc = Document(OUTPUT)

C_NAVY   = RGBColor(0x1F, 0x35, 0x64)
C_TEAL   = RGBColor(0x00, 0x70, 0x70)
C_AMBER  = RGBColor(0xC5, 0x5A, 0x11)
C_GREEN  = RGBColor(0x37, 0x86, 0x30)
C_LGREY  = RGBColor(0xF2, 0xF2, 0xF2)
C_DGREY  = RGBColor(0x40, 0x40, 0x40)
C_RED    = RGBColor(0xC0, 0x00, 0x00)
C_PURPLE = RGBColor(0x5B, 0x2C, 0x6F)
C_DKGRN  = RGBColor(0x14, 0x5A, 0x32)

def set_font(run, bold=False, italic=False, size=11, color=C_DGREY):
    run.bold = bold; run.italic = italic
    run.font.size = Pt(size); run.font.color.rgb = color

def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18); p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text); set_font(run, bold=True, size=18, color=C_NAVY)

def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text); set_font(run, bold=True, size=14, color=C_TEAL)

def h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text); set_font(run, bold=True, size=12, color=C_NAVY)

def h4(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text); set_font(run, bold=True, size=11, color=C_PURPLE)

def body(text, indent=0):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(indent * 0.25)
    run = p.add_run(text); set_font(run, size=10.5)

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text); set_font(run, size=10.5)

def note(label, text, color=C_AMBER):
    p = doc.add_paragraph()
    r1 = p.add_run(f"{label}: "); set_font(r1, bold=True, size=10.5, color=color)
    r2 = p.add_run(text); set_font(r2, size=10.5)

def code(text):
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text); set_font(run, bold=True, size=9.5, color=C_NAVY)

def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(6)
    run = p.add_run("─" * 85); set_font(run, size=8, color=RGBColor(0xCC, 0xCC, 0xCC))

def shade_cell(cell, rgb):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}")
    shd.set(qn('w:val'), 'clear'); tcPr.append(shd)

def simple_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]; shade_cell(cell, C_NAVY)
        p = cell.paragraphs[0]; run = p.add_run(h)
        set_font(run, bold=True, size=10, color=RGBColor(0xFF,0xFF,0xFF))
    for ri, row_data in enumerate(rows):
        row = t.add_row()
        bg = C_LGREY if ri % 2 == 0 else RGBColor(0xFF,0xFF,0xFF)
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]; shade_cell(cell, bg)
            p = cell.paragraphs[0]; run = p.add_run(str(val))
            set_font(run, size=10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t

# ── BULL MARKET PART HEADER PAGE ────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(40)
run = p.add_run("PART II — BULL MARKET TRADING")
set_font(run, bold=True, size=24, color=C_DKGRN)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Weinstein Commander Suite v4.5  ·  Stage 2 Advance Strategies")
set_font(run, italic=True, size=13, color=C_TEAL)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    "Strategy v4.5  ·  Dashboard v63.3  ·  Ultimate Screener v3.6\n"
    "Beta Edition v2.3  ·  Fundamental X-Ray v2.2  ·  Swing Zigzag v6.0"
)
set_font(run, size=11, color=C_DGREY)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 18 — PHILOSOPHY & THE UNIFIED METHOD
# ═══════════════════════════════════════════════════════════════════════════
h1("Section 18 — Bull Market Philosophy: The Unified Method")

h2("18.1 The Foundation: Why Two Masters?")
body(
    "The Commander Suite is built on the intersection of two independently validated methodologies: "
    "Stan Weinstein's Stage Analysis (1988) and Mark Minervini's Trend Template (SEPA, 2013). "
    "Neither system alone is sufficient — Weinstein filters the market to only uptrending stocks, "
    "Minervini filters those to only the strongest. Trading the overlap eliminates 90% of the universe "
    "and focuses attention on the institutional-grade setups that generate outsized returns."
)

h2("18.2 Stan Weinstein's Stage Analysis")
body(
    "Weinstein's framework classifies every stock into one of four lifecycle stages based on its "
    "30-Week SMA (Simple Moving Average) direction and price position relative to that MA. "
    "The core principle: institutions move in cycles. You must align with the institutional cycle, "
    "not fight it."
)

simple_table(
    ["Stage", "Name", "30W SMA Direction", "Price vs 30W SMA", "Volume Pattern", "Action"],
    [
        ["1", "Basing / Accumulation", "Flat", "Oscillating around MA", "Contracting (drying up)", "Watch only — base building"],
        ["2", "Advancing / Markup", "Rising", "Above MA (and rising)", "Expanding on up-days", "BUY ZONE — this is where wealth is created"],
        ["3", "Topping / Distribution", "Flattening/Rolling", "Choppy around MA", "Heavy, two-sided churn", "Exit — institutions are distributing"],
        ["4", "Declining / Markdown", "Falling", "Below MA (and falling)", "Expanding on down-days", "NEVER buy — capital destruction zone"],
    ],
    col_widths=[0.5, 1.6, 1.6, 1.8, 1.7, 1.8]
)
note("CORE RULE", "Only buy in Stage 2. Never average down in Stage 4. "
     "The 30-Week SMA is the single most important line on any chart.", color=C_DKGRN)

h2("18.3 Mark Minervini's Trend Template")
body(
    "Within Stage 2, Minervini's Trend Template applies six simultaneous filters to identify stocks "
    "with exceptional structural momentum. All six must be TRUE for the template to pass."
)
simple_table(
    ["#", "Criterion", "Condition", "Why It Matters"],
    [
        ["1", "Price above 50-DMA", "Close > SMA50", "Short-term trend is up"],
        ["2", "Price above 150-DMA", "Close > SMA150", "Medium-term trend is up"],
        ["3", "Price above 200-DMA", "Close > SMA200", "Long-term trend is up (institutional baseline)"],
        ["4", "SMA50 above SMA150", "SMA50 > SMA150", "Short > medium = aligned uptrend"],
        ["5", "SMA150 above SMA200", "SMA150 > SMA200", "Medium > long = full stack alignment"],
        ["6", "Price proximity to high", "Close ≥ 75% of 52W high AND ≥ 30% above 52W low", "Stock is strong, not just recovering from deep correction"],
    ],
    col_widths=[0.4, 2.2, 2.0, 2.4]
)

h2("18.4 The Unified Principle")
p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.3)
run = p.add_run(
    '"Stage 2 filters the universe to only uptrending stocks. '
    "Minervini's template filters that universe to only the strongest. "
    'You then trade the catalyst within that elite group."'
)
set_font(run, italic=True, bold=True, size=12, color=C_NAVY)
body("")

body(
    "The result: from an NSE/BSE universe of 5,000+ stocks, you typically have 30–80 stocks "
    "passing both filters in a bull market. Your job becomes selecting the best 2–6 of those "
    "based on catalyst quality, fundamental strength, and available capital slots."
)

h2("18.5 The Three-Layer Architecture")
body("Every bull market trade goes through three validation layers in sequence:")
simple_table(
    ["Layer", "Tool", "Question Answered", "Minimum Threshold"],
    [
        ["1. Universe Filter", "Beta Screener v2.3 / Ultimate v3.6", "Is this stock technically eligible?", "Stage 2 + Score ≥ 40 + Catalyst > 0"],
        ["2. Quality Gate", "Fundamental X-Ray v2.2", "Does the business deserve capital?", "Overall Grade ≥ B (score ≥ 9/17)"],
        ["3. Structure & Execution", "Zigzag v6.0 + Strategy v4.5 + Dashboard v63.3", "Is the entry precise and risk-defined?", "Uptrend HH/HL + Choppiness ≤ 4 + Clear Room to T1"],
    ],
    col_widths=[1.5, 2.0, 2.3, 2.2]
)

h2("18.6 Bull Market vs Recovery: The Key Differences")
simple_table(
    ["Parameter", "Bull Market", "Recovery (Phase 1)"],
    [
        ["Stage Required", "Stage 2 (all 6 Minervini criteria)", "Stage 1 Constructive or early Stage 2"],
        ["RS Requirement", "Leading or Weakening (positive)", "Just positive (RS > 0)"],
        ["Chandelier Multiplier", "3.0× ATR14 (tighter trail)", "3.5× ATR14 (wider for volatility)"],
        ["T1 Target", "2.5R (bull) / 2.0R (correction)", "2.0R (all regimes)"],
        ["T2 Target", "3.5R (bull) / 3.0R (bear)", "2.5R maximum"],
        ["Positional Time Stop", "6 weeks", "Applied via Time Stop input in Dashboard"],
        ["Swing Time Stop", "10 days", "10 days"],
        ["Max Slots", "6 (Bull) / 4 (Correction)", "2–4 (regime-dependent)"],
        ["Fundamental Minimum", "Grade B (positional) / Grade C (swing)", "RFF ≥ 1 check (lighter filter)"],
        ["Kelly Multiplier", "Up to 1.25× (A+ setups)", "Standard 1.0×"],
    ],
    col_widths=[2.5, 2.5, 2.5]
)

divider()

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 19 — THE SIX BULL MARKET CATALYSTS
# ═══════════════════════════════════════════════════════════════════════════
h1("Section 19 — The Six Bull Market Catalysts: Complete Playbooks")

h2("19.1 Catalyst Overview")
body(
    "The Commander Suite detects six distinct entry catalysts, each corresponding to a different "
    "market behaviour pattern. They are not interchangeable — each has its own win rate profile, "
    "holding period, target structure, and risk characteristics. Know all six."
)
simple_table(
    ["ID", "Code", "Name", "Type", "Win Rate Profile", "Best Regime"],
    [
        ["2", "POS-BO", "Positional Breakout", "Trend Following", "Highest in Bull markets", "Stage 2 full confluence"],
        ["3", "SWG-PB", "Swing Pullback", "Mean Reversion in Trend", "High / best R:R", "Stage 2, any regime"],
        ["1", "POS-ACCUM", "Positional Accumulation", "Early Positioning", "High if patient", "Stage 1/2 with OBV divergence"],
        ["4", "SWG-BO", "Swing VCP Breakout", "Momentum Breakout", "Good in Bull markets", "VCP structure in Stage 2"],
        ["6", "SWG-GAP", "Gap & Go", "Momentum", "Variable / time-sensitive", "News/earnings gap in Stage 2"],
        ["5", "SWG-REV", "Mean Reversion", "Counter-swing", "Lowest / use sparingly", "Strong Stage 2 only, small size"],
    ],
    col_widths=[0.4, 1.0, 1.8, 1.8, 1.8, 1.8]
)

h2("19.2 POS-BO — Positional Breakout (The A-Grade Setup)")
body(
    "The POS-BO is the flagship setup of the Commander Suite — a Stage 2 breakout above a "
    "multi-week base with institutional volume confirmation. These trades are held for weeks "
    "to months with a trailing stop (Chandelier Exit), targeting 2.5R and beyond."
)

h3("19.2.1 Required Conditions (All Must Be True)")
simple_table(
    ["Condition", "Requirement", "Why"],
    [
        ["Stage", "Weekly state machine = Stage 2 (SMA150 rising, price above SMA150)", "Institutional cycle is up"],
        ["Trend Aligned", "Daily close > SMA50 > SMA150 > SMA200", "Full Minervini stack"],
        ["Trend Template", "Close within 25% of 52W high AND ≥70% above 52W low", "Not buying deep in a hole"],
        ["RS Filter", "Mansfield RS > 0 vs both Nifty 50 and CNX500", "Stock leading the market"],
        ["Sector Stage", "Sector in Stage 1 or Stage 2", "Sector tide is supporting"],
        ["Vol Accumulation", "10+ days with vol > avg AND close in top 60% of range", "Institutional pre-loading"],
        ["MA Squeeze", "|EMA20 − SMA50| / SMA50 ≤ 6% (squeeze threshold)", "Price coiled, energy building"],
        ["BB Squeeze", "Bollinger Band width < 100-day average BB width", "Volatility at minimum before explosion"],
        ["Stage 2 Duration", "≤ 20 weeks in Stage 2 (fresh, not extended)", "Early in the move = most runway"],
        ["Alpha Score", "≥ 60 out of 100 (quality gate)", "Minimum institutional backing"],
    ],
    col_widths=[2.2, 2.8, 2.0]
)

h3("19.2.2 Trigger Conditions (Entry Signal Fires)")
simple_table(
    ["Trigger", "Condition"],
    [
        ["Breakout Level", "Close > highest(high, 20 bars)[1] — breaks 20-day pivot high"],
        ["Volume Expansion", "Volume > SMA50(volume) × 1.5 — institutional conviction required"],
        ["Price Action", "Close > Open AND (close − low) > (high − close) — bullish close in upper range"],
        ["Extension Check", "(close − breakout level) / breakout level ≤ 5% — not chasing"],
        ["Confirmation Mode", "Daily Close (not intraday touch) — strict mode reduces false signals"],
    ],
    col_widths=[2.5, 4.5]
)

h3("19.2.3 Trade Management")
bullet("Initial Stop: Structural low (10-bar low) minus 0.2× ATR14")
bullet("T1 (2.5R): Sell 30% of position — move remainder stop to breakeven")
bullet("Runner: Trail with Chandelier Exit (3.0× ATR14 in bull market, 3.5× in bear)")
bullet("CE tightens to 2.25× at 1.5R profit, tightens to 1.5× at 3.0R profit")
bullet("Stage 4 exit: Immediate close if weekly state machine detects Stage 4")
bullet("Time Stop: 6 weeks (30 bars) with < 0.5R profit → exit (capital velocity rule)")
note("RULE", "POS-BO is a position trade. Do not micro-manage it daily. "
     "Set the Chandelier stop, let the algorithm trail it, and only act on signals.", color=C_DKGRN)

h2("19.3 SWG-PB — Swing Pullback (The Best Risk:Reward Setup)")
body(
    "The SWG-PB is the highest-probability, best risk:reward setup in the suite. "
    "It buys Stage 2 stocks that have pulled back to the 20-EMA with volume dry-up — "
    "the Minervini 'healthy pullback' concept. These trades are held 1–3 weeks."
)

h3("19.3.1 Required Conditions")
simple_table(
    ["Condition", "Requirement"],
    [
        ["Stage / Trend", "Close > SMA150 AND SMA150 > SMA200 AND SMA50 > SMA150 (partial Minervini stack)"],
        ["Momentum Pocket", "Daily RSI ≥ 40 AND ≤ 60 (not overbought, not broken)"],
        ["Weekly RSI Floor", "Weekly RSI > 60 (macro momentum still intact)"],
        ["EMA Zone", "Price within EMA20 ± 1.5% (in the pullback strike zone)"],
        ["Pullback Depth", "Depth from recent swing high ≤ 15% (healthy pullback, not breakdown)"],
        ["Volume Dry-Up", "On down-day bars: volume < SMA50(volume) × 0.70 (institutions not selling)"],
        ["Reversal Pattern", "Bullish Engulfing OR Pin Bar/Hammer OR Morning Star on entry bar"],
        ["Micro Edge", "Close > CPR (Central Pivot Range) AND Close > Monthly VWAP AND MA squeeze active"],
        ["Alpha Score", "≥ 60 (quality gate)"],
    ],
    col_widths=[2.5, 4.5]
)

h3("19.3.2 Trade Management")
bullet("Initial Stop: EMA20 × (1 − 1%) — just below the EMA that the trade is bouncing from")
bullet("T1: 2.5R in bull market / 2.0R in correction — sell 50% of position")
bullet("Move stop to breakeven after T1 hit")
bullet("T2: 3.5R in bull / 3.0R in bear — sell 50% of remaining (25% of original)")
bullet("Final Runner (25%): Trail with 20-EMA until daily close breaks below EMA20")
bullet("50MA Fail Exit: If daily close < SMA50 → momentum thesis broken, close immediately")
bullet("Time Stop: 10 days with < 0.5R profit → exit (do not let it become dead money)")

h2("19.4 POS-ACCUM — Smart Money Accumulation (The Early Entry)")
body(
    "POS-ACCUM is an early-stage positional setup where OBV (On-Balance Volume) is making "
    "new highs while price is still lagging in a flat base. This is institutional accumulation "
    "before the breakout. Patient traders who enter here get the best prices."
)

h3("19.4.1 Required Conditions")
simple_table(
    ["Condition", "Requirement"],
    [
        ["Stage Gate", "Stage 1 Base OR Stage 2 Uptrend (not Stage 3 or 4)"],
        ["OBV Breakout", "OBV > highest(OBV, 50 bars)[1] — OBV making 50-bar highs while price lags"],
        ["OBV Trending", "OBV 13-bar linear regression rising (sustained, not single-bar spike)"],
        ["Price Lagging", "Close < highest(high, 50 bars)[1] × 0.90 — price at least 10% below 50-bar high"],
        ["Flat Base", "StdDev(close, 50) < SMA50 × 10% — price in tight consolidation"],
        ["RS Filter", "Mansfield RS > 0 vs CNX500"],
        ["Sector Stage", "Sector in Stage 1 or Stage 2 (FIX v4.5)"],
        ["Macro Edge", "Vol accumulation confirmed (10+ days, 2+ closes with V > V_avg, top 60% range)"],
        ["Alpha Score", "≥ 60"],
    ],
    col_widths=[2.5, 4.5]
)

h3("19.4.2 Trade Management")
bullet("Initial Stop: Structural 10-bar low minus 0.2× ATR14")
bullet("T1: 2.5R — sell 30% of position, trail remainder with Chandelier Exit")
bullet("This is a positional trade — same trailing/exit logic as POS-BO once in profit")
bullet("Patience required: may take 2–8 weeks before price catches up to OBV signal")
note("NOTE", "POS-ACCUM is an OBV divergence trade. The stock may not immediately move. "
     "That is expected. The value is in buying before the crowd recognises the accumulation.", color=C_AMBER)

h2("19.5 SWG-BO — Swing VCP Breakout (The Momentum Breakout)")
body(
    "SWG-BO targets a Volatility Contraction Pattern (VCP) — a series of progressively tighter "
    "consolidations before a pivot break. Minervini's canonical setup. Requires the Anti-Algo Gate "
    "(close in top 40% of bar) to filter out pump-and-dump spikes."
)

h3("19.5.1 Required Conditions")
simple_table(
    ["Condition", "Requirement"],
    [
        ["Minervini Trend", "Close > SMA50 > SMA150 > SMA200 (full Minervini stack)"],
        ["VCP Tightness", "ATR10 < SMA50(ATR10) × 1.5 — volatility contracting below normal"],
        ["Pivot Break", "Close > highest(high, 15 bars)[1] — breaks 15-bar pivot high"],
        ["Volume Expansion", "Volume > SMA50(volume) × 1.5 — buying power behind the break"],
        ["Anti-Algo Gate", "Close in top 40% of bar's range: (close−low)/(high−low) ≥ 0.60"],
        ["Alpha Score", "≥ 60 (FIX v4.5 — this gate was missing in v4.4)"],
    ],
    col_widths=[2.5, 4.5]
)

h3("19.5.2 Trade Management")
bullet("Initial Stop: Lowest(low, 5 bars) minus 0.2% padding")
bullet("T1: 2.5R (bull) / 2.0R (bear) — sell 50% of position")
bullet("T2: 3.5R (bull) / 3.0R (bear) — sell 50% of remaining (25% of original)")
bullet("Final Runner (25%): Trail with 20-EMA close")
bullet("50MA Fail Exit: Close below SMA50 → exit immediately")
note("CRITICAL", "The Anti-Algo Gate is non-negotiable. A SWG-BO signal where the close "
     "is in the bottom 40% of the bar is a HFT trap. If the gate fails, DO NOT ENTER.", color=C_RED)

h2("19.6 SWG-GAP — Gap & Go (The Institutional Momentum Play)")
body(
    "SWG-GAP captures institutional gap-ups (≥4%) driven by earnings, news, or macro catalysts. "
    "These are high-gain, volatile, time-sensitive trades. The entry is tight (gap bar low is the stop), "
    "the gain can be large if the gap holds."
)

h3("19.6.1 Required Conditions")
simple_table(
    ["Condition", "Requirement"],
    [
        ["Gap Size", "(open − high[1]) / high[1] × 100 ≥ 4.0% — minimum gap threshold"],
        ["Volume", "Volume > SMA50(volume) × 3.0 — three times normal (institutional conviction)"],
        ["Close Position", "Close > Open (up-close on gap bar — not a reversal gap)"],
        ["Intraday Position", "(close − low) / (high − low) ≥ 0.60 — close in top 40% of gap bar (anti-algo, FIX v4.5)"],
        ["OBV Trend", "OBV 13-bar linear regression rising before gap (institutional pre-loading, FIX v3.6)"],
        ["Stage Gate", "Stage 2 OR Stage 1 (not Stage 3 or 4 — FIX v4.5)"],
        ["Market Health", "Market confirmed healthy (not BEARISH)"],
    ],
    col_widths=[2.5, 4.5]
)

h3("19.6.2 Trade Management")
bullet("Initial Stop: Low of the gap bar minus 0.2% padding (TIGHT stop)")
bullet("T1: 2.5R — sell 50% of position at first target")
bullet("T2: 3.5R — sell 50% of remaining")
bullet("Runner (25%): Trail with 20-EMA")
bullet("Time Stop: 10 days with < 0.5R — gaps that fail to follow through are dead money")
note("RULE", "Do NOT buy a gap the next day if you missed it. Gap trades are entry-day plays. "
     "If you miss the gap bar, wait for a SWG-PB setup on the same stock.", color=C_AMBER)

h2("19.7 SWG-REV — Mean Reversion (The Speculative Bounce)")
body(
    "SWG-REV is the weakest setup in the suite — lowest win rate, most dependent on market regime. "
    "It targets extreme RSI-3 oversold conditions (< 20) in an otherwise healthy Stage 2 stock. "
    "Use only with reduced size and only in confirmed bull markets."
)

h3("19.7.1 Required Conditions (All Must Pass — Stricter than Other Setups)")
simple_table(
    ["Condition", "Requirement"],
    [
        ["RSI Oversold", "RSI(3) < 20 OR RSI(14) < 30 — extreme short-term oversold"],
        ["Weekly Floor", "Weekly RSI > 40 — macro momentum not broken"],
        ["Stage Gate", "Stage 1 Base OR Stage 2 Uptrend (NOT Stage 3 or 4 — FIX v4.5)"],
        ["Market Health", "Market confirmed healthy (not BEARISH)"],
        ["Pattern Trigger", "Any reversal candle: Engulfing, Hammer, or Morning Star"],
        ["Volume Spike", "Volume > SMA50(volume) × 1.5 on entry bar"],
        ["EMA Extension", "(EMA20 − close) / EMA20 > 5% — price stretched well below EMA20"],
    ],
    col_widths=[2.5, 4.5]
)

h3("19.7.2 Trade Management")
bullet("Initial Stop: Lowest(low, 5 bars) minus padding")
bullet("T1: 2.0R only (always — no dynamic adjustment for SWG-REV)")
bullet("Time Stop: Hard 5-day exit regardless of P&L (mean reversion must resolve quickly)")
bullet("Position Size: Use 50%–75% of normal size (lower conviction setup)")
note("WARNING", "SWG-REV was fixed in v4.5 to require Stage 1 or Stage 2. "
     "In prior versions it fired in Stage 4, causing catastrophic dead-cat bounce entries. "
     "If your system version is v4.4 or earlier, manually verify Stage before taking this setup.", color=C_RED)

h2("19.8 The SMC Liquidity Sweep Enhancement")
body(
    "Any catalyst tagged with '(SMC)' has additionally confirmed an institutional Liquidity Sweep: "
    "the 50-DMA was swept (price briefly dipped below it intrabar) then reclaimed with volume surge. "
    "This is Smart Money Concepts (SMC) confirmation that institutional traders absorbed retail stop-outs."
)
simple_table(
    ["SMC Condition", "Requirement"],
    [
        ["5-bar Low", "lowest(low, 5 bars) < SMA50 — stop-sweep below 50-DMA"],
        ["Recovery", "Close > SMA50 — price reclaimed the level"],
        ["Volume Confirmation", "Volume > SMA50(volume) × 1.5 — institutional absorption confirmed"],
        ["Score Boost", "+20 points added to Alpha Score — hyper-conviction signal"],
    ],
    col_widths=[2.5, 4.5]
)
note("GUIDE", "An SMC tag on any catalyst means treat it as a higher-conviction setup. "
     "Apply the Kelly 1.25× multiplier and prioritise these over non-SMC signals.", color=C_DKGRN)

divider()

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 20 — STRATEGY v4.5 COMPLETE REFERENCE
# ═══════════════════════════════════════════════════════════════════════════
h1("Section 20 — Weinstein-Minervini Strategy v4.5: Complete Reference")

h2("20.1 Purpose")
body(
    "The Strategy v4.5 is the mechanical execution engine of the Commander Suite. "
    "It serves three roles simultaneously: (1) a live signal generator showing buy/sell signals "
    "on the chart, (2) a backtesting engine for validating the system on any instrument, and "
    "(3) a dynamic position sizing calculator via the Risk Allocator table."
)
body("Apply on Daily timeframe. Load alongside Dashboard v63.3 and Zigzag v6.0 for complete analysis.")

h2("20.2 All Input Fields — Complete Reference")

h3("Group 1: Master Configuration")
simple_table(
    ["Input Name", "Default", "Options", "Description"],
    [
        ["Strategy Mode", "Hybrid (Both)", "Breakout Only / Pullback Only / Hybrid (Both)", "Which catalyst types fire entries. Hybrid trades all 6 types."],
        ["Asset Type", "Auto", "Auto / ETF / Stock", "Drives risk %. Auto detects via syminfo.type. ETFs get 1.0%, Stocks get 0.75%."],
    ],
    col_widths=[2.2, 1.2, 2.5, 2.1]
)

h3("Group 2: Portfolio & Risk")
simple_table(
    ["Input Name", "Default", "Description"],
    [
        ["Total Portfolio Capital", "100,000", "Your actual trading capital. All position sizes are calculated from this."],
        ["ETF Risk %", "1.0%", "Max loss per ETF trade as % of portfolio. ETFs are less volatile so wider risk allowed."],
        ["Stock Risk %", "0.75%", "Max loss per stock trade. Standard equity risk per trade."],
        ["Max Allocation per Trade", "25,000", "Hard capital cap per single position. Prevents over-concentration."],
        ["Max Open Positions", "6", "Portfolio heat cap. System will not generate new signals if 6 positions are open."],
        ["Positional Time Stop (Weeks)", "6", "Exit positional trade if < 0.5R profit after 6 weeks. Capital velocity rule."],
        ["Swing Time Stop (Days)", "10", "Exit swing trade if < 0.5R profit after 10 days. Dead money prevention."],
        ["Kelly Fraction Sizing", "ON", "Scales risk up (1.25×) for A+ setups, down (0.75×) for weak setups."],
    ],
    col_widths=[2.5, 1.0, 3.5]
)

h3("Group 3: Strategy A — Weinstein Positional")
simple_table(
    ["Input Name", "Default", "Description"],
    [
        ["[EDGE] Stage 2 Breakout", "ON", "Enables POS-BO catalyst. Main positional entry."],
        ["[EDGE] Smart Money Accum (OBV)", "ON", "Enables POS-ACCUM catalyst. OBV divergence entries."],
        ["Breakout Lookback (Days)", "20", "Pivot high period. Close > highest(high, 20)[1] = breakout."],
        ["Max Stage 2 Duration (Weeks)", "20", "Skip entry if stock has been in Stage 2 > 20 weeks. Avoids extended late-stage buys."],
        ["Max Breakout/Gap Extension (%)", "5.0%", "Reject entry if price already ≥5% above breakout level. Prevents chasing."],
        ["Breakout Confirmation", "Daily Close", "Daily Close = strict (must close above). High/Touch = aggressive (any intraday touch)."],
        ["Global ATR Length", "14", "ATR period for all volatility calculations and stop placements."],
    ],
    col_widths=[2.5, 1.0, 3.5]
)

h3("Group 4: Mathematical Edges (Backtested)")
simple_table(
    ["Edge", "Default", "Effect When Enabled"],
    [
        ["Macro Edge: Institutional Vol Bias", "ON",
         "Requires VWMA(50) > SMA(50) vol shelf AND 10+ accumulation days before firing STRONG BUY. "
         "Penalises setups lacking institutional volume participation. If OFF: −20 alpha score penalty."],
        ["Micro Edge: Price & Squeeze Validation", "ON",
         "SWG-PB only fires if: Price > CPR AND Price > Monthly VWAP AND MA squeeze active (|EMA20−SMA50|/SMA50 ≤ 6%). "
         "Eliminates pullbacks in low-quality environments."],
        ["SMC Liquidity Sweep Filter", "ON",
         "When 50MA swept intrabar then reclaimed with volume: +20 alpha score boost, (SMC) tag added to signal label. "
         "Hyper-conviction institutional footprint confirmation."],
    ],
    col_widths=[2.8, 0.8, 3.4]
)
note("BEST PRACTICE", "Leave all three edges ON at all times unless you are deliberately "
     "stress-testing a setup without edge filters. The backtests show a 15–25% improvement "
     "in win rate when all three are active vs all three disabled.", color=C_DKGRN)

h3("Group 5: Strategy B — Hybrid Swing Engine")
simple_table(
    ["Input Name", "Default", "Description"],
    [
        ["[EDGE] Pullback / Moving Avg Bounce", "ON", "Enables SWG-PB. EMA20 bounce entries."],
        ["[EDGE] Momentum Breakout (VCP)", "ON", "Enables SWG-BO. VCP pivot break entries."],
        ["[EDGE] Mean Reversion (Oversold)", "ON", "Enables SWG-REV. RSI-3 < 20 bounce entries. Use sparingly."],
        ["[EDGE] Gap & Go", "ON", "Enables SWG-GAP. ≥4% gap-up institutional entries."],
        ["Daily RSI Pullback Pocket Min/Max", "40 / 60", "SWG-PB RSI range. Price must be in momentum zone, not overbought."],
        ["Weekly RSI Momentum Min", "60", "Weekly RSI must exceed 60 for SWG-PB to fire. Macro momentum gate."],
        ["Daily EMA Length", "20", "The anchor moving average for pullback zone calculation."],
        ["EMA Zone Tolerance (%)", "1.5%", "Price must be within EMA20 ± 1.5% to be in pullback zone."],
        ["Max Pullback Depth (%)", "15%", "Pullback from swing high must be ≤15%. Deeper = broken, not a pullback."],
        ["Trigger: Bullish Engulfing", "ON", "Allow bullish engulfing candle as SWG-PB entry trigger."],
        ["Trigger: Pin Bar/Hammer", "ON", "Allow hammer/pin bar as SWG-PB entry trigger."],
        ["Trigger: Morning Star", "ON", "Allow morning star 3-candle pattern as SWG-PB trigger."],
        ["Use Dynamic R:R Targets", "ON", "T1 = 2.5R in bull market (market health=true), 2.0R in bear."],
        ["Pullback Target 1 R:R (Static)", "2.0", "Used only when Dynamic R:R is OFF."],
        ["Pullback SL Padding %", "0.2%", "Small buffer below structural low for initial stop."],
        ["Pullback Trail Buffer %", "1.0%", "EMA20 trailing stop buffer: trail at EMA20 × (1−1%)."],
    ],
    col_widths=[2.8, 0.8, 3.4]
)

h3("Group 6: Shared Filters")
simple_table(
    ["Input Name", "Default", "Description"],
    [
        ["Use Relative Strength Filter", "ON", "Requires Mansfield RS > 0 vs Nifty 500 AND sector. Blocks entries on laggards."],
        ["Auto-Detect Sector", "ON", "Automatically maps stock to sector index via syminfo.industry keyword matching."],
        ["Manual Sector Override", "NSE:BANKNIFTY", "Used only if Auto-Detect is OFF or incorrectly classifying the stock."],
    ],
    col_widths=[2.5, 1.2, 3.3]
)

h3("Group 7: Core Mechanics (Keep Synced with Dashboard v63.3)")
simple_table(
    ["Input Name", "Default", "Description"],
    [
        ["Est. Pullback Lookback (Days)", "21", "Fibonacci swing window. Used for Fib retracement levels and pullback depth calculation."],
        ["Trailing Stop ATR Multiplier (Base)", "3.0", "Chandelier Exit base multiplier. 3.0× ATR in bull market. Adaptive: tightens as profit grows."],
        ["Adaptive TSL: Tighten Trigger (R)", "1.5", "At 1.5R profit: CE multiplier scales from 3.0× down toward 2.25×."],
        ["Adaptive TSL: Tightening Ratio", "0.75", "New CE multiplier = base × 0.75 when 1.5R trigger hit."],
        ["Adaptive TSL: Super-Tight Trigger (R)", "3.0", "At 3.0R profit: CE multiplier scales down further toward 1.5×."],
        ["Adaptive TSL: Super-Tight Ratio", "0.50", "New CE multiplier = base × 0.50 when 3.0R trigger hit."],
    ],
    col_widths=[2.8, 0.8, 3.4]
)

h3("Group 8: Risk Allocator (Manual Input)")
simple_table(
    ["Input Name", "Default", "Description"],
    [
        ["Show Allocator Table", "ON", "Toggle the on-chart position sizing table (bottom-right by default)."],
        ["Entry Price", "0.0", "Your planned entry price. Set with confirm dialog. If 0 and a live signal exists, uses signal price."],
        ["Stop Price", "0.0", "Your planned stop loss. Set with confirm dialog. Required for quantity calculation."],
        ["Lot Multiplier", "1", "For F&O (e.g., 50 for Nifty lot). Keep at 1 for all equity trades."],
        ["Table Position", "bottom_right", "Position of the allocator table on chart."],
    ],
    col_widths=[2.5, 1.0, 3.5]
)

h2("20.3 Reading the Risk Allocator Table")
body(
    "The Risk Allocator appears as a compact table on the chart. It auto-populates from "
    "three sources: (1) your manual Entry/Stop inputs, (2) the live algorithm signal if no manual "
    "inputs are set, or (3) the open trade if one is active."
)
simple_table(
    ["Row", "Content", "Key Notes"],
    [
        ["Asset Type", "ETF / Stock", "Drives the risk % selection"],
        ["Entry", "Planned entry price", "Orange text"],
        ["Stop", "Planned stop price", "Red text"],
        ["Adjusted Risk %", "Dynamic risk after volatility + Kelly", "Shows the actual % being risked, post-adjustment"],
        ["QUANTITY", "Shares to buy", "YELLOW — the critical number. Use this exactly."],
        ["Invested", "Total capital deployed", "Entry × Quantity"],
        ["Risk Amount", "Rupees at stake", "Should be ~0.75% of portfolio for stocks"],
        ["Target 1 (xR)", "First profit target price", "Green — 2.5R in bull, 2.0R in bear"],
        ["Target 2 (xR)", "Runner profit target price", "Lime — 3.5R in bull, 3.0R in bear"],
        ["Trade Validation", "Clear Room to T1 OR T1 > Overhead Resistance", "ORANGE WARNING if blocked — skip the trade"],
    ],
    col_widths=[2.0, 2.5, 2.5]
)

h2("20.4 Kelly Fraction Sizing Logic")
body("The Kelly system dynamically adjusts your risk percentage based on setup quality:")
simple_table(
    ["Conditions Met", "Kelly Multiple", "Effective Risk (Stock)", "Interpretation"],
    [
        ["Stage 2 + RS improving + Vol > avg (all 3)", "1.25×", "0.9375% (~0.94%)", "A+ setup — bet bigger"],
        ["Two of three conditions above", "1.0×", "0.75%", "Standard setup — normal size"],
        ["One or zero conditions met", "0.75×", "0.5625% (~0.56%)", "Weak setup — bet smaller"],
    ],
    col_widths=[3.0, 1.2, 1.5, 2.3]
)
body("Kelly is further modified by:")
bullet("Market Regime Discount: Bear market → base_risk − 0.25% (floor: 0.25%)")
bullet("Volatility Discount: vol_discount = 3.0% / (ATR14/close×100) — capped at 1.0×")
bullet("Final: adjusted_risk = regime_risk × min(vol_discount × kelly_mult, 1.0)")
note("NOTE", "The combined floor is 0.25% — the system never risks less than a quarter percent "
     "even in the most adverse conditions.", color=C_AMBER)

h2("20.5 Adaptive Chandelier Exit (Trailing Stop)")
body(
    "The Chandelier Exit (CE) is the primary trailing stop for positional trades. "
    "It ratchets upward as price rises, never steps backward."
)
code("CE = Highest High (ATR_length bars) − (ATR14 × ce_multiplier)")
body("The multiplier adapts as the trade profits:")
simple_table(
    ["Profit Level", "CE Multiplier", "Effect"],
    [
        ["Entry → 1.5R", "3.0× ATR (bull) / 3.5× ATR (bear)", "Maximum room — trade needs to breathe"],
        ["1.5R → 3.0R", "3.0 × 0.75 = 2.25× ATR", "Tightening — protecting accumulated profits"],
        ["3.0R+", "3.0 × 0.50 = 1.5× ATR", "Super-tight — holding gains aggressively"],
    ],
    col_widths=[1.8, 2.5, 2.7]
)
note("RULE", "Never manually move the Chandelier stop down. It only ratchets up. "
     "If you are tempted to widen the stop to 'give the trade more room,' that is a losing behavior.", color=C_RED)

h2("20.6 Complete Exit Logic Reference")
h3("Positional (POS-BO, POS-ACCUM) Exits")
simple_table(
    ["Exit Type", "Trigger", "Action"],
    [
        ["T1 Partial (30%)", "Price hits entry + initial_risk × 2.5R", "Sell 30% of position. is_breakeven = TRUE (SL moves to entry)."],
        ["Chandelier Trail", "Daily close < Chandelier Exit level", "Close remaining 70% at market next open."],
        ["Stage 4 Emergency", "Weekly state machine transitions to Stage 4", "Immediate full exit regardless of CE level."],
        ["Time-Decay", "6 weeks (30 bars) in trade with < 0.5R profit", "Close full position. Capital velocity rule — no exceptions."],
        ["Bear Market CE", "Market health turns bearish", "CE multiplier widens to 3.5× — prevents whipsaw exits."],
    ],
    col_widths=[2.0, 2.8, 2.2]
)

h3("Swing (SWG-*) Exits")
simple_table(
    ["Exit Type", "Trigger", "Action"],
    [
        ["T1 Partial (50%)", "Price hits entry + risk × 2.5R (bull) / 2.0R (bear)", "Sell 50% of position. Stop moves to breakeven."],
        ["T2 Exit (25%)", "Price hits entry + risk × 3.5R (bull) / 3.0R (bear)", "Sell 50% of remaining position (25% of original)."],
        ["20-EMA Trail", "Daily close < EMA20 × (1−1%)", "Close final runner (remaining 25%)."],
        ["50MA Fail", "Daily close < SMA50", "Close full position immediately. Momentum thesis broken."],
        ["Time-Decay", "10 days with < 0.5R profit", "Close full position."],
        ["SWG-REV Time Stop", "5 days hard stop (always)", "Mean reversion must complete within 5 days or the thesis is wrong."],
    ],
    col_widths=[2.0, 2.8, 2.2]
)

divider()

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 21 — ULTIMATE SCREENER v3.6 COMPLETE REFERENCE
# ═══════════════════════════════════════════════════════════════════════════
h1("Section 21 — Commander Screener Ultimate v3.6: Complete Reference")

h2("21.1 Purpose")
body(
    "The Ultimate Screener v3.6 is a simultaneous multi-stock batch analysis tool that ranks up to "
    "30 stocks by your chosen criterion on a single chart view. It is the first tool in the daily "
    "workflow — converting a large watchlist into a ranked shortlist of actionable candidates."
)
body("Apply on any chart (e.g., NSE:NIFTY). The underlying chart symbol is irrelevant.")

h2("21.2 Input Fields")
simple_table(
    ["Input Name", "Default", "Description"],
    [
        ["Watchlist Symbols", "(comma-separated)", "Up to 30 NSE/BSE tickers, e.g., 'NSE:RELIANCE, NSE:TCS'. Parser trims spaces automatically."],
        ["Benchmark (For RS)", "NSE:NIFTY", "Benchmark used for Mansfield RS calculation for all 30 stocks."],
        ["Rank/Sort By", "Score/Grade", "Primary sort criterion. Options: Score/Grade, Volume Thrust, Proximity to 52-Wk High, Proximity to 50-SMA, Relative Strength (Qtr)."],
        ["Sort Direction", "Descending", "Best-to-worst (highest score first) or Worst-to-best."],
        ["Table Position", "bottom_center", "9 position options for table placement on chart."],
    ],
    col_widths=[2.5, 1.5, 3.0]
)

h2("21.3 The 17-Column Table — Complete Column Reference")
body("The table renders a Market Health Banner, a header row, then one row per stock.")
body("Banner format: '⚔ COMMANDER SCREENER | Market: [state] | Bench: [symbol]'")

simple_table(
    ["Col #", "Header", "Values / Range", "Colour Guide", "How to Read"],
    [
        ["0", "RANK", "1–30", "Green = S2 UP, Gold = S1, Orange = S3, Red = S4", "Position after sorting"],
        ["1", "TICKER", "Stock name", "White text", "NSE: prefix removed for readability"],
        ["2", "SCORE", "Stars + number + Grade", "Lime ≥80, Yellow ≥60, Orange ≥40, Red <40", "Core quality metric — focus on 4+ stars"],
        ["3", "STAGE", "STAGE 1–4 + variant", "Green = S2, Amber = S1, Red = S4", "Must be Stage 2 for entries"],
        ["4", "STG WKS", "e.g., '12W', '5W'", "White", "Weeks in current stage. <20W = fresh, >20W = extended"],
        ["5", "RS / SEC RS", "LEADING / Sec: BEAT", "Teal = leading, Red = lagging", "2-line: market RS + sector RS"],
        ["6", "50D SLOPE", "Rising ↗ / Flat / Falling ↘", "Green = rising, Red = falling", "50-DMA trend direction"],
        ["7", "VOLATILITY", "SQUEEZE / Normal / Expanded", "Green = squeeze, Red = expanded", "SQUEEZE = ideal entry zone"],
        ["8", "VOL%", "e.g., '185%'", "Green ≥150%, Amber 120–150%", "Volume vs 50-day average"],
        ["9", "ATH DIST%", "e.g., '-3.5%'", "Green = within 5% of ATH", "Distance from 52-week high"],
        ["10", "50SMA DIST%", "e.g., '+8.2%'", "Green = above 50SMA", "Distance from 50-DMA"],
        ["11", "RSI", "e.g., '58.3' or '72.1 (OB)'", "Red = overbought >70", "RSI(14). OB = overbought"],
        ["12", "TRADE", "T1:xxx | BLOCKED | AVOID | WAIT", "Green = active, Red = blocked", "If BLOCKED: skip. If T1:xxx = entry is valid."],
        ["13", "CATALYST", "POS-BO / SWG-PB / etc.", "Colour by catalyst type", "The specific setup that triggered"],
        ["14", "STYLE", "BOTH / POS / SWING / WAIT", "Gold = both, Blue = positional, Fuchsia = swing", "Trade horizon recommendation"],
        ["15", "PERSONA", "LEADER / MOMENTUM / VOLATILE / etc.", "Lime = leader, Red = volatile", "Stock character — avoid VOLATILE for new entries"],
        ["16", "PRICE", "Current close", "White", "Live price for reference"],
    ],
    col_widths=[0.5, 1.2, 1.8, 1.8, 2.2]
)

h2("21.4 Scoring Engine — Complete Breakdown (0–100)")
simple_table(
    ["Component", "Points Added", "Condition"],
    [
        ["Stage 2 (UP)", "+25", "State machine confirms Stage 2: SMA150 rising + price above SMA150"],
        ["Stage 1 (BASE)", "+10", "Stage 1 base confirmed (above SMA200, not yet Stage 2)"],
        ["RS Leading", "+15", "Mansfield RS slope > 0.02 — stock clearly outperforming benchmark"],
        ["RS Improving", "+5", "RS slope neutral but RS value positive — recovery in progress"],
        ["50DMA Rising + Aligned", "+10", "Slope > threshold AND SMA50 > SMA150 > SMA200 — perfect stack"],
        ["50DMA Rising Only", "+5", "Slope > threshold but SMA stack not fully aligned"],
        ["VCP Squeeze", "+10", "ATR10 < SMA50(ATR10) × 1.5 — classic Minervini VCP tight base"],
        ["Volume Thrust", "+10", "Volume > SMA50(volume) × 1.5 on an up-close bar"],
        ["RSI 50–70", "+5", "RSI in healthy trending zone (not overbought, not broken)"],
        ["SMC Liquidity Sweep", "+20", "5-bar low swept SMA50, reclaimed with volume (v3.6 ENHANCEMENT)"],
        ["Volume Shelf", "+5", "VWMA(50) > SMA(50) — institutional volume participation (v3.3 ENHANCEMENT)"],
        ["Alpha Negative Penalty", "−10", "Q3 or H1 alpha vs benchmark is negative — stock underperforming badly"],
    ],
    col_widths=[2.5, 1.0, 3.5]
)

h2("21.5 Grade-to-Action Mapping")
simple_table(
    ["Score Range", "Stars", "Grade", "Trade Action"],
    [
        ["80–100", "★★★★★", "A", "STRONG BUY — Full position, all edges confirmed"],
        ["60–79", "★★★★☆", "B", "BUY — Good setup, standard position"],
        ["40–59", "★★★☆☆", "C", "WATCHLIST — Stage 1 base or partial confluence. Monitor."],
        ["20–39", "★★☆☆☆", "D", "AVOID — Trend broken or RS lagging badly"],
        ["0–19", "★☆☆☆☆", "F", "AVOID / MONITOR SHORT — Stage 4 or fully broken"],
    ],
    col_widths=[1.2, 1.0, 0.8, 4.0]
)

h2("21.6 TRADE Column: ACTIVE vs BLOCKED Explained")
body("The TRADE column is critical — it tells you whether the setup has clear room to reach T1.")
simple_table(
    ["Value", "Meaning", "Action"],
    [
        ["T1: xxx.xx | SL: xxx.xx", "Setup valid AND T1 target has clear overhead path", "Enter the trade — risk/reward is valid"],
        ["BLOCKED (T1=xxx.xx)", "T1 > 52W high or major overhead resistance — poor RR", "Skip this setup. Wait for new base above resistance."],
        ["AVOID", "Stage 4 OR Grade D — do not trade", "Remove from watchlist unless monitoring for reversal"],
        ["WAIT", "Stage 2 but no active catalyst detected", "Add to watchlist. Set BoS alert on Zigzag."],
    ],
    col_widths=[2.5, 2.5, 2.0]
)

h2("21.7 Stage State Machine (Hysteresis-Aware, v3.6)")
body(
    "The stage classifier in v3.6 uses a stateful machine with hysteresis — it remembers the "
    "previous stage and only transitions when conditions are persistently met. This eliminates "
    "the false stage-flipping from single-bar noise that plagued earlier versions."
)
body("Weekly slope proxy formula (v3.6 FIX — aligned with Strategy v4.4):")
code("w_slope = (SMA150 − SMA150[20]) / 20.0")
body("State transitions require:")
bullet("Stage 1 → Stage 2: w_slope > dynThresh AND close > SMA150")
bullet("Stage 2 → Stage 3: NOT (w_slope > dynThresh) AND NOT (close > SMA150)")
bullet("Stage 3 → Stage 4: w_slope < −dynThresh AND close < SMA150")
bullet("Stage 4 → Stage 1: NOT (w_slope < −dynThresh) AND close > SMA150")
bullet("dynThresh = SMA150 × 0.0005 (dynamic threshold scales with price level)")

h2("21.8 Sector RS Detection (v3.3 Enhancement)")
body("The Ultimate Screener pre-fetches all 8 NSE sector indices and calculates sector RS for each stock:")
simple_table(
    ["Sector Code", "Sector", "Index Used"],
    [
        ["1", "Banking & Finance", "NSE:BANKNIFTY"],
        ["2", "Information Technology", "NSE:CNXIT"],
        ["3", "Pharmaceuticals & Healthcare", "NSE:CNXPHARMA"],
        ["4", "Automobiles", "NSE:CNXAUTO"],
        ["5", "FMCG & Consumer", "NSE:CNXFMCG"],
        ["6", "Metals & Mining", "NSE:CNXMETAL"],
        ["7", "Energy (Oil, Gas, Power)", "NSE:CNXENERGY"],
        ["8", "Real Estate", "NSE:CNXREALTY"],
        ["0", "Default / Other", "Uses benchmark (NSE:NIFTY)"],
    ],
    col_widths=[1.2, 2.5, 2.0]
)
body("Sector RS result codes in the SEC RS column:")
bullet("🔥 BEAT — stock outperforming its sector (strongest signal)")
bullet("Neutral+ — slight outperformance vs sector")
bullet("Neutral- — slight underperformance vs sector")
bullet("LAGGING — clear underperformer within sector (avoid)")

h2("21.9 Sorting Engine")
simple_table(
    ["Sort By Option", "Sort Value Used", "Best For"],
    [
        ["Score/Grade", "score + (volThrust / 1000)", "General daily scan — finds best overall setups"],
        ["Volume Thrust", "volThrust (V/Vavg × 100)", "Finding stocks with institutional buying today"],
        ["Proximity to 52-Wk High", "distFromHigh × −1 (closest first)", "Finding stocks near breakout levels"],
        ["Proximity to 50-SMA", "|distFrom50| × −1 (closest first)", "Finding pullback candidates near 50-DMA"],
        ["Relative Strength (Qtr)", "2=LEADING, 1=Improving, 0=other", "Finding RS leaders regardless of price action"],
    ],
    col_widths=[2.5, 2.2, 2.3]
)

doc.add_page_break()
doc.save(OUTPUT)
print(f"Part 5 saved: {OUTPUT}")
print("Sections 18–21 complete.")
