"""
Recovery Trading Manual — Generator Part 3
Section 16 — Dashboard v63.3 Portfolio Management
Section 17 — Dashboard v63.3 All Input Fields
Appends to existing document.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = r"C:\Users\jayra\Documents\GeminiVSCode\Recovery_Trading_Manual.docx"
doc = Document(OUTPUT)

C_NAVY  = RGBColor(0x1F, 0x35, 0x64)
C_TEAL  = RGBColor(0x00, 0x70, 0x70)
C_AMBER = RGBColor(0xC5, 0x5A, 0x11)
C_GREEN = RGBColor(0x37, 0x86, 0x30)
C_LGREY = RGBColor(0xF2, 0xF2, 0xF2)
C_DGREY = RGBColor(0x40, 0x40, 0x40)
C_RED   = RGBColor(0xC0, 0x00, 0x00)

def set_font(run, bold=False, italic=False, size=11, color=C_DGREY):
    run.bold = bold; run.italic = italic
    run.font.size = Pt(size); run.font.color.rgb = color

def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18); p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text); set_font(run, bold=True, size=18, color=C_NAVY)

def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text); set_font(run, bold=True, size=14, color=C_TEAL)

def h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text); set_font(run, bold=True, size=12, color=C_NAVY)

def body(text, indent=0):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(indent * 0.25)
    run = p.add_run(text); set_font(run, size=10.5)

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text); set_font(run, size=10.5)

def note(label, text, color=C_AMBER):
    p = doc.add_paragraph()
    r1 = p.add_run(f"{label}: "); set_font(r1, bold=True, size=10.5, color=color)
    r2 = p.add_run(text); set_font(r2, size=10.5)

def code(text):
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text); set_font(run, bold=True, size=10, color=C_NAVY)

def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(6)
    run = p.add_run("─" * 85); set_font(run, size=8, color=RGBColor(0xCC, 0xCC, 0xCC))

def shade_cell(cell, rgb):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}")
    shd.set(qn('w:val'), 'clear'); tcPr.append(shd)

def simple_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]; shade_cell(cell, C_NAVY)
        p = cell.paragraphs[0]; run = p.add_run(h)
        set_font(run, bold=True, size=10, color=RGBColor(0xFF,0xFF,0xFF))
    for ri, row_data in enumerate(rows):
        row = t.add_row()
        bg = C_LGREY if ri % 2 == 0 else RGBColor(0xFF,0xFF,0xFF)
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]; shade_cell(cell, bg)
            p = cell.paragraphs[0]; run = p.add_run(str(val))
            set_font(run, size=10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 16 — DASHBOARD v63.3 PORTFOLIO MANAGEMENT
# ═══════════════════════════════════════════════════════════════════════════
h1("Section 16 — Dashboard v63.3: Portfolio Management Features")

h2("16.1 Overview")
body("Dashboard v63.3 is a 4-column × 50-row Pine Script table displayed at "
     "position.bottom_right (default). It integrates real-time technical analysis "
     "with portfolio management fields to give a single-screen view of any open position.")
body("Colour defaults: background=white, text=black, value text=black (light theme). "
     "All colours are user-configurable via inputs.")

h2("16.2 Left Column — Full Row Order")
body("The left column runs in this exact order top-to-bottom:")

h3("Block 1: Trade Summary")
simple_table(
    ["Row Label", "Content / Logic"],
    [
        ["RECOMMENDATION", "Primary action: BUY / HOLD / TRIM / EXIT / WATCH"],
        ["CATALYST (Edge)", "Active catalyst tag: POS-ACCUM / POS-BO / SWG-GAP / SWG-BO / SWG-PB / SWG-REV / NONE"],
        ["ASSET QUALITY", "AI grade from ai_grading_engine: 5★ / 4★ / 3★ / 2★ / 1★"],
        ["ACTION SIGNAL", "7-component star rating (Trend/Mom/Vol/VP/PA/ADX/LiqSweep) shown as ★★★☆☆☆☆"],
        ["TRADE STYLE", "Positional / Swing — set by user input p_style"],
        ["PERSONA", "Aggressive / Balanced / Conservative — affects SL width and target"],
        ["PORTFOLIO HEALTH", "Composite portfolio risk score from open positions"],
    ],
    col_widths=[2.5, 4.5]
)

h3("Block 2: Active Trade Rows (Conditional — only when p_entry > 0)")
simple_table(
    ["Row Label", "Content / Logic"],
    [
        ["MY TRADE", "Shows entry price vs current price, direction arrow"],
        ["ENTRY DATE", "Date when p_entry was set (user input)"],
        ["DAYS HELD", "Calendar days since entry date"],
        ["Time Stop", "Maximum hold duration; triggers exit signal when reached"],
        ["Dynamic Profile", "Current trade stage: Early / Developing / Extended / At Risk"],
        ["Current R:R", "Live risk:reward ratio based on current price vs entry and SL"],
    ],
    col_widths=[2.5, 4.5]
)
note("NOTE", "These 6 rows are hidden when p_entry = 0 (no active trade set).")

h3("Block 3: Institutional Activity & Filters (separator row above)")
simple_table(
    ["Row Label", "Content / Logic"],
    [
        ["Daily Close > CPR", "Pass/Fail — daily close above Central Pivot Range"],
        ["Price > M-VWAP", "Pass/Fail — price above monthly VWAP"],
        ["Vol Shelf (VWMA)", "Pass/Fail — volume-weighted MA supporting price"],
        ["Vol Accum (Setup)", "Pass/Fail — 10-day accumulation check (vol_acc_days=10, min ratio=vol_acc_min=2)"],
        ["VCP Tightness", "Pass/Fail — ATR10 < ATR40 × 1.5 (vcp_thresh=1.5 hidden constant)"],
        ["Anti-Algo BO Gate", "Pass/Fail — close in top 60% of range + up close required for SWG-BO catalyst"],
    ],
    col_widths=[2.5, 4.5]
)

h3("Block 4: Breadth & Sector (separator row above)")
simple_table(
    ["Row Label", "Content / Logic"],
    [
        ["Mkt Health (N500)", "BULLISH (Confirmed) / CORRECTION (In Uptrend) / RECOVERY (Weak) / BEARISH"],
        ["Sector Info", "Sector name and current stage"],
        ["Sector Velocity (ROC)", "🔥 ACCELERATING LEADER / 🛑 EXHAUSTED LEADER / ⚡ HIDDEN ACCUMULATION / 🧊 DEAD MONEY / 🧊 Static"],
        ["RS (vs Nifty 50)", "Mansfield RS value vs ^NSEI"],
        ["RS (vs N500)", "Mansfield RS value vs ^CRSLDX — primary RS measure"],
        ["RS (sector)", "Stock rank within sector by RS"],
        ["Sector Stage (W)", "Weekly stage of the sector: Stage 1/2/3/4"],
        ["Next Earnings", "Days until next earnings announcement"],
    ],
    col_widths=[2.5, 4.5]
)

h2("16.3 Right Column — Full Row Order")

h3("Block 1: Positional Analysis")
simple_table(
    ["Row Label", "Content / Logic"],
    [
        ["POSITIONAL (header)", "Section divider"],
        ["Market Structure", "Overall market structure: Uptrend / Downtrend / Ranging"],
        ["Master Trend (W)", "Weekly trend direction from higher-timeframe analysis"],
        ["30-Week MA Slope", "Rising / Flat / Falling — core Weinstein stage indicator"],
        ["Price > 30W MA", "Yes / No — Weinstein primary condition"],
        ["Overhead Resist", "Nearest overhead resistance level and % distance"],
        ["Room for Trade", "Yes / No — checks if T1 target is not blocked by resistance"],
        ["52W High Space", "% distance from current price to 52-week high"],
        ["ATH Space", "% distance from current price to all-time high"],
        ["Weekly PB Health", "HEALTHY (Low Vol) / NORMAL / WEAK (Selling)"],
    ],
    col_widths=[2.5, 4.5]
)

h3("Block 2: Swing Analysis")
simple_table(
    ["Row Label", "Content / Logic"],
    [
        ["SWING (header)", "Section divider"],
        ["Daily Trend (D)", "Daily trend direction: Uptrend / Downtrend / Consolidation"],
        ["Trend Template (>200d)", "Pass/Fail — Minervini Trend Template: 8 criteria check"],
        ["Price > 50 DMA", "Yes / No — daily momentum check"],
        ["50 DMA Slope", "Rising / Flat / Falling"],
        ["Price Action Structure", "Higher Highs/Lows / Lower Highs/Lows / Ranging"],
        ["Volume Trend", "Increasing / Decreasing / Neutral on recent swing"],
        ["Volatility State (ADR)", "VOLATILITY SQUEEZE (Ready) / VCP (True Dry Up) / VCP (Actionable) / Vol Expansion / Normal Volatility"],
        ["EMA Proximity", "RESISTANCE (Hit Head) / SUPPORT (Bounced) / NEAR (x.x%) / PRICE ABOVE / PRICE BELOW"],
        ["S/R Proximity", "Nearest support/resistance level and % distance"],
        ["RSI & Divergence", "RSI14 value + bullish/bearish divergence flag if detected"],
        ["Daily Pullback Health", "HEALTHY (Low Vol) / NORMAL / WEAK (Selling)"],
    ],
    col_widths=[2.5, 4.5]
)

h3("Block 3: Alpha Screener")
simple_table(
    ["Row Label", "Content / Logic"],
    [
        ["ALPHA SCREENER (header)", "Section divider"],
        ["Top 5 Slots", "Symbols 1–5 ranked by Alpha Score. Min score: 40 (hidden constant). Shows symbol + score."],
    ],
    col_widths=[2.5, 4.5]
)
note("NOTE", "Alpha Screener only populates when p_watch_list input contains symbols. "
     "It re-ranks on every bar close.")

h2("16.4 TRADE STATUS Display Logic")
body("Full conditional logic for the TRADE STATUS label (appears in RECOMMENDATION row when p_entry > 0):")
code("if close < effectiveSL:")
code("    Status = '⚠️ STOP HIT (P&L%)'")
code("else if curPnL > 0:")
code("    cond_trim = dailyRsi > 75 OR absDist > 15.0 OR earnings < 3 days")
code("    cond_add  = dTrendDir==1 AND dailyRsi<70 AND close>EMA20 AND NOT isSqueeze")
code("    if cond_trim: Status = 'TRIM (P&L%)'")
code("    elif cond_add: Status = 'SCALE IN (P&L%)'")
code("    else: Status = 'HOLD (P&L%)'")
code("else:")
code("    Status = 'HOLD (P&L%)'  // losing but above stop")
body("Panic Exit override:")
code("if close < wMA40 (40W MA): Status = 'EXIT NOW (40WMA Break)'")

h2("16.5 Hidden Internal Constants")
body("These values are hardcoded in the Pine source and cannot be changed from the UI:")
simple_table(
    ["Constant", "Value", "Purpose"],
    [
        ["vol_acc_days", "10", "Number of days used for volume accumulation check"],
        ["vol_acc_min", "2", "Minimum vol/avg ratio to pass accumulation check"],
        ["vol_avg_len", "50", "Lookback for average volume calculation"],
        ["vcp_thresh", "1.5", "ATR ratio threshold for VCP detection"],
        ["vdu_pct", "0.30", "Volume dry-up threshold: relVol < 0.30× average"],
    ],
    col_widths=[2.0, 1.0, 4.0]
)

divider()

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 17 — DASHBOARD v63.3 ALL INPUT FIELDS
# ═══════════════════════════════════════════════════════════════════════════
h1("Section 17 — Dashboard v63.3: All Input Fields Reference")

h2("17.1 Display & Appearance")
simple_table(
    ["Input Name", "Type", "Default", "Description"],
    [
        ["Show Dashboard", "bool", "true", "Toggle entire table on/off"],
        ["Table Position", "string", "bottom_right", "position.bottom_right — change via dropdown"],
        ["Background Color", "color", "white", "Table background (light theme default)"],
        ["Text Color", "color", "black", "Row label text color"],
        ["Value Text Color", "color", "black", "Row value text color"],
        ["Bullish Color", "color", "green", "Color for bullish readings"],
        ["Bearish Color", "color", "red", "Color for bearish readings"],
        ["Neutral Color", "color", "gray", "Color for neutral readings"],
    ],
    col_widths=[2.2, 0.8, 1.2, 3.0]
)

h2("17.2 Trade Parameters")
simple_table(
    ["Input Name", "Type", "Default", "Description"],
    [
        ["p_entry", "float", "0.0", "Entry price. Set to 0 if no active trade. Unlocks trade rows."],
        ["p_sl", "price", "0.0", "Stop-loss level. Uses input.price() — gives draggable handle on chart."],
        ["p_entry_date", "string", "", "Entry date in YYYY-MM-DD format for Days Held calculation."],
        ["p_style", "string", "Swing", "Trade style: Positional or Swing — affects SL/target formulas."],
        ["p_persona", "string", "Balanced", "Risk persona: Aggressive / Balanced / Conservative."],
        ["p_time_stop", "int", "20", "Maximum trading days to hold before time stop triggers."],
    ],
    col_widths=[2.2, 0.8, 1.2, 3.0]
)
note("IMPORTANT", "p_sl uses input.price() (not input.float()). This creates a draggable "
     "orange horizontal line on the chart. Drag it to update — no need to type a value.")

h2("17.3 Benchmark & Sector")
simple_table(
    ["Input Name", "Type", "Default", "Description"],
    [
        ["Benchmark Symbol", "string", "NSE:NIFTY", "Primary benchmark for RS calculations. Usually NSE:NIFTY or ^CRSLDX."],
        ["Sector Symbol", "string", "", "Sector index for sector RS comparison (e.g., NSE:NIFTYIT)."],
        ["RS Lookback (weeks)", "int", "26", "Weeks for Mansfield RS SMA calculation."],
        ["Market Benchmark (N500)", "string", "^CRSLDX", "CNX500 symbol for market health calculation."],
    ],
    col_widths=[2.5, 0.8, 1.2, 2.7]
)

h2("17.4 Algo Signal Mode")
simple_table(
    ["Input Name", "Type", "Default", "Description"],
    [
        ["Enable Algo Signal", "bool", "false", "When true, auto-draws SL/T1/T2 lines when catalyst fires."],
        ["Algo SL Type", "string", "Positional", "Positional: d_l20 − (d_atr10 × 0.2). Swing: d_l20 × 0.998."],
        ["Show Algo Lines", "bool", "true", "Toggle visibility of auto-drawn SL/T1/T2 lines."],
        ["Algo Line Color", "color", "orange", "Color for algo-drawn lines."],
    ],
    col_widths=[2.2, 0.8, 1.4, 2.8]
)

h2("17.5 Alpha Screener")
simple_table(
    ["Input Name", "Type", "Default", "Description"],
    [
        ["Watch List", "string", "", "Comma-separated symbols to score (e.g., 'RELIANCE,TCS,INFY'). Max 20 recommended."],
        ["Min Score Display", "int", "40", "Only show symbols scoring ≥ this value in the top-5 list."],
        ["Max Slots", "int", "5", "Number of top symbols to display in Alpha Screener block."],
    ],
    col_widths=[2.2, 0.8, 1.0, 3.2]
)

h2("17.6 Technical Parameters")
simple_table(
    ["Input Name", "Type", "Default", "Description"],
    [
        ["ATR Length", "int", "14", "ATR lookback for volatility and SL calculations."],
        ["RSI Length", "int", "14", "RSI lookback period."],
        ["EMA Length", "int", "20", "Short-term EMA for proximity and trend checks."],
        ["Volume MA Length", "int", "50", "SMA period for average volume (used in relative vol calc)."],
        ["30W MA Length", "int", "30", "Weeks for the primary Weinstein SMA (maps to ~150 trading days)."],
        ["ADX Length", "int", "14", "ADX/DI lookback for trend strength check in Action Signal."],
        ["Chandelier Mult", "float", "3.5", "ATR multiplier for Chandelier Exit SL. Recovery uses 3.5×."],
        ["Squeeze ATR Ratio", "float", "0.60", "ATR5/ATR20 threshold for VOLATILITY SQUEEZE state."],
    ],
    col_widths=[2.5, 0.7, 1.0, 3.0]
)

h2("17.7 Stage Analysis Overrides")
simple_table(
    ["Input Name", "Type", "Default", "Description"],
    [
        ["SMA50 Length", "int", "50", "50-day SMA for Minervini Trend Template and stage checks."],
        ["SMA150 Length", "int", "150", "150-day SMA for full SMA stack alignment."],
        ["SMA200 Length", "int", "200", "200-day SMA for primary stage boundary."],
        ["Weekly SMA Length", "int", "30", "30-week SMA (Weinstein's primary tool). Weekly timeframe."],
        ["Stage Override", "string", "Auto", "Force a stage display: Auto / Stage1 / Stage2 / Stage3 / Stage4."],
    ],
    col_widths=[2.5, 0.7, 1.0, 3.0]
)
note("NOTE", "Stage Override is for testing only. Leave on Auto for live trading.")

h2("17.8 Earnings Data")
simple_table(
    ["Input Name", "Type", "Default", "Description"],
    [
        ["Next Earnings Date", "string", "", "YYYY-MM-DD format. Dashboard calculates days until earnings."],
        ["Earnings Warning Days", "int", "3", "Days before earnings that trigger TRIM warning in Trade Status."],
    ],
    col_widths=[2.5, 0.7, 1.0, 3.0]
)

doc.add_page_break()
doc.save(OUTPUT)
print(f"Part 3 saved: {OUTPUT}")
print("Sections 16–17 complete.")
