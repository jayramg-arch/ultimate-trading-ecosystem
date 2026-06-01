"""
Recovery Trading Manual — Generator Part 4
Appendix A: Module Reference (A.1–A.4)
Appends to existing document.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = r"C:\Users\jayra\Documents\GeminiVSCode\Recovery_Trading_Manual.docx"
doc = Document(OUTPUT)

C_NAVY  = RGBColor(0x1F, 0x35, 0x64)
C_TEAL  = RGBColor(0x00, 0x70, 0x70)
C_AMBER = RGBColor(0xC5, 0x5A, 0x11)
C_GREEN = RGBColor(0x37, 0x86, 0x30)
C_LGREY = RGBColor(0xF2, 0xF2, 0xF2)
C_DGREY = RGBColor(0x40, 0x40, 0x40)
C_RED   = RGBColor(0xC0, 0x00, 0x00)

def set_font(run, bold=False, italic=False, size=11, color=C_DGREY):
    run.bold = bold; run.italic = italic
    run.font.size = Pt(size); run.font.color.rgb = color

def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18); p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text); set_font(run, bold=True, size=18, color=C_NAVY)

def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text); set_font(run, bold=True, size=14, color=C_TEAL)

def h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text); set_font(run, bold=True, size=12, color=C_NAVY)

def body(text, indent=0):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(indent * 0.25)
    run = p.add_run(text); set_font(run, size=10.5)

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text); set_font(run, size=10.5)

def note(label, text, color=C_AMBER):
    p = doc.add_paragraph()
    r1 = p.add_run(f"{label}: "); set_font(r1, bold=True, size=10.5, color=color)
    r2 = p.add_run(text); set_font(r2, size=10.5)

def code(text):
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text); set_font(run, bold=True, size=9.5, color=C_NAVY)

def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(6)
    run = p.add_run("─" * 85); set_font(run, size=8, color=RGBColor(0xCC, 0xCC, 0xCC))

def shade_cell(cell, rgb):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}")
    shd.set(qn('w:val'), 'clear'); tcPr.append(shd)

def simple_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]; shade_cell(cell, C_NAVY)
        p = cell.paragraphs[0]; run = p.add_run(h)
        set_font(run, bold=True, size=10, color=RGBColor(0xFF,0xFF,0xFF))
    for ri, row_data in enumerate(rows):
        row = t.add_row()
        bg = C_LGREY if ri % 2 == 0 else RGBColor(0xFF,0xFF,0xFF)
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]; shade_cell(cell, bg)
            p = cell.paragraphs[0]; run = p.add_run(str(val))
            set_font(run, size=10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t

# ═══════════════════════════════════════════════════════════════════════════
# APPENDIX A — MODULE REFERENCE
# ═══════════════════════════════════════════════════════════════════════════
h1("Appendix A — Python Module Reference")

body("This appendix documents the four Python modules that power the screening, "
     "grading, and risk management workflows. All files are located in: "
     "C:\\Users\\jayra\\Documents\\GeminiVSCode\\")

divider()

# A.1 — chartink_scanner_pro.py
h2("A.1 — chartink_scanner_pro.py")
body("Version: v2.0 (redesigned April 2026). "
     "Purpose: Pulls candidate lists from Chartink using 7 pre-defined scanner configurations. "
     "Scanners 5, 6, and 7 are the Recovery phase scanners.")

h3("Recovery Scanners (5, 6, 7)")
simple_table(
    ["Scanner #", "Edge", "Output File", "Key Criteria"],
    [
        ["5", "REV-RS", "Recovery_RS_Survivors.csv",
         "RS Line > 30W SMA of RS Line + above SMA200 + above SMA50 + weekly RSI > 55 + close > ₹100"],
        ["6", "REV-CB", "Recovery_Climax_Bounce.csv",
         "RS Line > 30W SMA + price ≥5% below SMA200 + RSI < 55 + close > ₹100"],
        ["7", "REV-EARLY", "Recovery_Early_Birds.csv",
         "SMA50 ≥ 92% of SMA200 + above SMA50 and SMA150 + RS positive + weekly RSI > 50 + close > ₹100"],
    ],
    col_widths=[0.8, 1.0, 2.5, 2.9]
)

h3("Design Philosophy (v2.0)")
body("v2.0 redesign targets quality companies in temporary corrections, NOT distressed stocks. "
     "The ₹100 minimum price filter enforces liquidity. The RS positive filter ensures "
     "the stock is not a market laggard.")

h3("Usage")
code("python chartink_scanner_pro.py")
body("Scanners run sequentially. Output CSVs are written to the same directory. "
     "Uses Selenium for browser-based Chartink query execution — ensure ChromeDriver is installed.")

divider()

# A.2 — recovery_screener.py
h2("A.2 — recovery_screener.py")
body("Version: v1.1. "
     "Purpose: Deep-scoring pipeline that applies RFF, regime gate, and composite scoring "
     "to the Chartink output. Produces a ranked shortlist for human review.")

h3("Configuration Block (top of file — edit these)")
simple_table(
    ["Parameter", "Default", "Description"],
    [
        ["signal_hold_days", "5", "Days a signal remains visible after first detection (weekend buffer)"],
        ["cb_climax_window", "10", "REV-CB: climax candle must be within this many bars"],
        ["cb_stretch_pct", "8.0", "REV-CB: minimum % stretch below 30W SMA"],
        ["cb_vol_mult", "2.0", "REV-CB: minimum volume multiplier vs 50-day avg for climax bar"],
        ["rs_bo_len", "20", "REV-RS: bars used for breakout pivot high calculation"],
        ["vol_confirm_mult", "1.5", "Breakout volume must be ≥ this × 50-day avg"],
        ["early_pivot_len", "15", "REV-EARLY: bars used for pivot high"],
        ["near_gc_pct", "5.0", "REV-EARLY: SMA50 must be within this % of SMA200 from below"],
        ["vcp_atr_mult", "1.5", "REV-EARLY: ATR10 < ATR40 × this value for VCP detection"],
        ["mkt_correction_pct", "7.0", "Regime: CNX500 must be ≥ this % below 52W high for regime block"],
        ["min_stock_correction_pct", "10.0", "Regime: stock must be ≥ this % below 52W high for regime block"],
        ["rff_min_score", "1", "Minimum RFF checks passed for a stock to be considered"],
        ["data_lookback_days", "400", "Days of historical price data to download per symbol"],
        ["download_delay_sec", "0.4", "Delay between yfinance downloads (rate limit protection)"],
    ],
    col_widths=[2.5, 0.8, 3.9]
)

h3("Signal Priority Values")
simple_table(
    ["Signal Value", "Edge", "Priority"],
    [
        ["4", "REV-EARLY", "Highest — cleanest setups, first allocation"],
        ["3", "REV-RS", "High — confirmed RS leadership"],
        ["2", "REV-CB", "Medium — bounce setup, T1 only"],
        ["1", "CB-Watch", "Low — watchlist, not yet actionable"],
        ["0", "None", "No signal — excluded from output"],
    ],
    col_widths=[1.0, 1.5, 4.5]
)

h3("Composite Score Formula (0–20)")
body("Score = RFF(0–6) + RS_positive(+1) + RS_strong_≥1.0(+1) + "
     "discount_depth(+1/+2/+3) + regime(+1/+2) + stage(+1) + "
     "edge_fired(+1/+2/+3) + volume(+1/+2) + intersection_bonus(+1)")

h3("Output")
code("Output: Recovery_Screener_Results.csv")
body("Sorted: Signal descending, then Score descending. "
     "Columns: Symbol, Signal, Score, Edge, RFF_Score, RS, Catalyst, Regime, Stage, Notes.")

divider()

# A.3 — ai_grading_engine.py
h2("A.3 — ai_grading_engine.py")
body("Purpose: Hybrid Quant + AI narrative grading engine. "
     "Fetches real market data via yfinance, runs deterministic scoring, "
     "then uses an LLM to generate a one-line institutional justification.")

h3("Primary Function")
code("get_weinstein_score(symbol, sector, ltp, buy_price, rs_status='Unknown', stage='Stage 2')")
body("Returns a dict:")
simple_table(
    ["Key", "Type", "Description"],
    [
        ["rating", "str", "5-Star / 4-Star / 3-Star / 2-Star / 1-Star"],
        ["reason", "str", "15-word max LLM-generated justification"],
        ["quant_grade", "str", "A / B / C / D / F"],
        ["quant_score", "int", "0–100 numeric score"],
        ["breakdown", "dict", "Component scores: SMA Stack / EMA Proximity / RSI / Volume / 52W High"],
        ["catalyst", "str", "POS-ACCUM / POS-BO / SWG-GAP / SWG-BO / SWG-PB / SWG-REV / NONE"],
    ],
    col_widths=[1.5, 0.8, 4.9]
)

h3("Quant Scoring Breakdown (0–100)")
simple_table(
    ["Component", "Max", "Scoring Logic"],
    [
        ["SMA Stack", "30 pts", "Perfect (C>50>150>200): 30 | Pullback (50>150>200): 20 | C>200 messy: 10 | C<200: 0"],
        ["EMA Proximity", "15 pts", "< 3% from EMA20: 15 | < 8%: 10 | < 15%: 5 | > 15%: 0"],
        ["RSI Health", "20 pts", "55–75: 20 | 45–55: 12 | >75 overbought: 5 | <45: 0"],
        ["Volume", "15 pts", "Rel vol > 1.5×: 15 | > 1.0×: 10 | < 1.0×: 3"],
        ["52W High Distance", "20 pts", "Within 5%: 20 | Within 15%: 12 | Within 30%: 5 | >30% off: 0"],
    ],
    col_widths=[2.0, 0.8, 4.4]
)

h3("Grade Mapping")
simple_table(
    ["Score", "Grade", "Star Rating"],
    [
        ["80–100", "A", "5-Star"],
        ["60–79", "B", "4-Star"],
        ["40–59", "C", "3-Star"],
        ["20–39", "D", "2-Star"],
        ["0–19", "F", "1-Star"],
    ],
    col_widths=[1.2, 1.0, 1.5]
)

h3("Caching")
body("Results are cached via ai_cache_manager.py. "
     "Cache key: {type: 'weinstein_grade_v2', symbol: 'SYMBOL'}. "
     "Cache is checked before any network call — subsequent runs for the same symbol "
     "are instant. Clear cache to force fresh data.")

divider()

# A.4 — ai_risk_manager.py
h2("A.4 — ai_risk_manager.py")
body("Purpose: Portfolio-level risk utilities. Called by Commander Web and the screener "
     "to provide ATR calculations, market health, correlation analysis, and portfolio vitals.")

h3("Exported Functions")
simple_table(
    ["Function", "Returns", "Description"],
    [
        ["get_atr(symbol, period=14)", "float", "ATR in price units. Fetches 60 days of daily data. "
         "Symbol auto-appended with .NS for NSE stocks."],
        ["get_market_health(benchmark='^NSEI')", "(bool, ltp, sma200)",
         "bool=True means price > SMA200 (healthy). Ticker map: ^CNX500 → ^CRSLDX."],
        ["get_nifty_correlation(symbol, period='60d')", "float or 'N/A'",
         "Pearson correlation vs Nifty 50 over specified period. -1 to +1."],
        ["get_noise_risk_stats(df)", "(int, list)",
         "Returns count of noisy positions and list of symbols where SL < 1.5×ATR14."],
        ["get_portfolio_correlation_matrix(symbols)", "(corr_df, shadow_pairs, div_score)",
         "Full correlation matrix. Shadow pairs = r > 0.85. Diversity score 0–1."],
        ["get_adaptive_atr_multiplier(symbol)", "float (1.5 / 2.0 / 2.5)",
         "ADR% < 1.5% → 1.5×. 1.5–3% → 2.0×. > 3% → 2.5×."],
        ["calculate_portfolio_vitals(closed_df, total_capital, open_df, live_map)", "dict",
         "Full performance stats: sharpe, max_drawdown, calmar, expectancy, win_rate, "
         "avg_win, avg_loss, total_return_pct, unrealized_pnl."],
    ],
    col_widths=[3.0, 1.2, 3.0]
)

h3("calculate_portfolio_vitals() — Input Format")
body("closed_df: DataFrame with columns [symbol, entry_price, exit_price, shares, date_in, date_out]")
body("open_positions_df: DataFrame with columns [symbol, entry_price, shares, current_price]")
body("live_map: dict {symbol: current_price} — live prices for open position MTM")
body("total_capital: float — total trading capital in rupees")

h3("Noise Risk — What It Means")
body("A 'noisy' position has a stop-loss placed so close to the entry price that "
     "normal intraday volatility (1.5× ATR14) could trigger it accidentally. "
     "Flagged positions should be reviewed: either widen the stop or reduce size.")

divider()

# ═══════════════════════════════════════════════════════════════════════════
# APPENDIX B — QUICK REFERENCE CARD
# ═══════════════════════════════════════════════════════════════════════════
h1("Appendix B — Quick Reference Card")

h2("Entry Checklist (All 10 Gates)")
simple_table(
    ["#", "Gate", "Pass Condition"],
    [
        ["1", "Market Health", "Not BEARISH"],
        ["2", "Regime Gate", "Three-Way Gate not fully blocked"],
        ["3", "RFF", "≥ rff_min_score checks passed"],
        ["4", "Composite Score", "≥ 8"],
        ["5", "AI Grade", "B or A (≥ 60/100)"],
        ["6", "Stage", "Stage 1 Constructive or Stage 2"],
        ["7", "Action Signal", "≥ 3/7 stars"],
        ["8", "Room for Trade", "T1 not blocked"],
        ["9", "Earnings", "Not within 3 days"],
        ["10", "Slot available", "Portfolio not at max slots for regime"],
    ],
    col_widths=[0.4, 2.0, 4.3]
)

h2("Target Table by Edge")
simple_table(
    ["Edge", "Stop", "T1", "T2", "Scale In?"],
    [
        ["REV-CB", "Below climax low − 0.5×ATR", "2.0R (exit 100%)", "None", "No"],
        ["REV-RS", "Chandelier 3.5×ATR14", "2.0R (exit 50%)", "2.5R (exit 50%)", "Yes if SCALE IN shows"],
        ["REV-EARLY", "Below VCP low − 0.2×ATR", "2.0R (exit 50%)", "2.5R (exit 50%)", "Yes if SCALE IN shows"],
    ],
    col_widths=[1.5, 2.2, 1.5, 1.5, 0.8]
)

h2("Weekly Workflow Summary")
simple_table(
    ["Day", "Task", "Tools"],
    [
        ["Sunday AM", "Run Chartink scanners", "chartink_scanner_pro.py"],
        ["Sunday AM", "Run deep scoring on results", "recovery_screener.py"],
        ["Sunday PM", "Grade top 10–15 candidates", "ai_grading_engine.py"],
        ["Sunday PM", "Review Dashboard on shortlist", "TradingView + Dashboard v63.3"],
        ["Sunday PM", "Update all open position SLs", "Dashboard p_sl draggable input"],
        ["Sunday PM", "Check portfolio vitals", "ai_risk_manager.py → calculate_portfolio_vitals()"],
        ["Daily pre-mkt", "Check regime + open position status", "Dashboard Market Health row"],
        ["Daily pre-mkt", "Check for ⚠️ STOP HIT overnight", "Dashboard TRADE STATUS"],
        ["Intraday", "Monitor TRIM / SCALE IN signals", "Dashboard RECOMMENDATION row"],
    ],
    col_widths=[1.8, 2.8, 2.6]
)

h2("Signal Hold Window")
body("Recovery screener signals remain visible for 5 trading days (signal_hold_days=5). "
     "This means a signal generated on Monday is still valid through the following Monday. "
     "After 5 days without entry, re-run the screener to confirm the setup is still active.")

h2("Key Formulas")
p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.3)
lines = [
    "Mansfield RS     = (RS_Line / SMA26_of_RS_Line − 1) × 10",
    "Position Size    = (Capital × Risk%) / (Entry − Stop)",
    "R:R              = (Target − Entry) / (Entry − Stop)",
    "Chandelier Stop  = Highest High (14 bars) − 3.5 × ATR14",
    "Algo SL Pos.     = d_l20 − (d_atr10 × 0.2)",
    "Algo SL Swing    = d_l20 × 0.998",
    "Algo T1          = Entry + Risk × 2.0",
    "Algo T2          = Entry + Risk × 3.5",
]
for line in lines:
    code(line)

doc.add_page_break()
doc.save(OUTPUT)
print(f"Part 4 saved: {OUTPUT}")
print("Appendix A + B complete. Manual is finished.")
