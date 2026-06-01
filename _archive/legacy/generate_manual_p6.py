"""
Recovery Trading Manual — Generator Part 6
BULL MARKET SECTION: Sections 22–26
- Section 22: Fundamental X-Ray v2.2 Complete Reference
- Section 23: Swing Zigzag Strict v6.0 Complete Reference
- Section 24: Beta Screener v2.3 Complete Reference
- Section 25: Commander Web v3.0 Complete Reference
- Section 26: Dashboard v63.3 Complete Deep Dive (Bull Market Focus)
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = r"C:\Users\jayra\Documents\GeminiVSCode\Recovery_Trading_Manual.docx"
doc = Document(OUTPUT)

C_NAVY   = RGBColor(0x1F, 0x35, 0x64)
C_TEAL   = RGBColor(0x00, 0x70, 0x70)
C_AMBER  = RGBColor(0xC5, 0x5A, 0x11)
C_GREEN  = RGBColor(0x37, 0x86, 0x30)
C_LGREY  = RGBColor(0xF2, 0xF2, 0xF2)
C_DGREY  = RGBColor(0x40, 0x40, 0x40)
C_RED    = RGBColor(0xC0, 0x00, 0x00)
C_PURPLE = RGBColor(0x5B, 0x2C, 0x6F)
C_DKGRN  = RGBColor(0x14, 0x5A, 0x32)

def sf(run, bold=False, italic=False, size=11, color=C_DGREY):
    run.bold=bold; run.italic=italic; run.font.size=Pt(size); run.font.color.rgb=color

def h1(text):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(18); p.paragraph_format.space_after=Pt(6)
    r=p.add_run(text); sf(r,bold=True,size=18,color=C_NAVY)

def h2(text):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(12); p.paragraph_format.space_after=Pt(4)
    r=p.add_run(text); sf(r,bold=True,size=14,color=C_TEAL)

def h3(text):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(2)
    r=p.add_run(text); sf(r,bold=True,size=12,color=C_NAVY)

def h4(text):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(6); p.paragraph_format.space_after=Pt(2)
    r=p.add_run(text); sf(r,bold=True,size=11,color=C_PURPLE)

def body(text,indent=0):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(4)
    p.paragraph_format.left_indent=Inches(indent*0.25)
    r=p.add_run(text); sf(r,size=10.5)

def bullet(text,level=0):
    p=doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent=Inches(0.25+level*0.2); p.paragraph_format.space_after=Pt(2)
    r=p.add_run(text); sf(r,size=10.5)

def note(label,text,color=C_AMBER):
    p=doc.add_paragraph()
    r1=p.add_run(f"{label}: "); sf(r1,bold=True,size=10.5,color=color)
    r2=p.add_run(text); sf(r2,size=10.5)

def code(text):
    p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(0.3)
    r=p.add_run(text); sf(r,bold=True,size=9.5,color=C_NAVY)

def divider():
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(6); p.paragraph_format.space_after=Pt(6)
    r=p.add_run("─"*85); sf(r,size=8,color=RGBColor(0xCC,0xCC,0xCC))

def shade_cell(cell,rgb):
    tc=cell._tc; tcPr=tc.get_or_add_tcPr()
    shd=OxmlElement('w:shd')
    shd.set(qn('w:fill'),f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}")
    shd.set(qn('w:val'),'clear'); tcPr.append(shd)

def tbl(headers,rows,col_widths=None):
    t=doc.add_table(rows=1,cols=len(headers))
    t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    hdr=t.rows[0]
    for i,h in enumerate(headers):
        cell=hdr.cells[i]; shade_cell(cell,C_NAVY)
        p=cell.paragraphs[0]; r=p.add_run(h); sf(r,bold=True,size=10,color=RGBColor(0xFF,0xFF,0xFF))
    for ri,rd in enumerate(rows):
        row=t.add_row()
        bg=C_LGREY if ri%2==0 else RGBColor(0xFF,0xFF,0xFF)
        for ci,val in enumerate(rd):
            cell=row.cells[ci]; shade_cell(cell,bg)
            p=cell.paragraphs[0]; r=p.add_run(str(val)); sf(r,size=10)
    if col_widths:
        for i,w in enumerate(col_widths):
            for row in t.rows: row.cells[i].width=Inches(w)
    doc.add_paragraph()
    return t

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 22 — FUNDAMENTAL X-RAY v2.2
# ═══════════════════════════════════════════════════════════════════════════
h1("Section 22 — Weinstein Fundamental X-Ray v2.2: Complete Reference")

h2("22.1 Purpose")
body(
    "The Fundamental X-Ray v2.2 is the quality gate of the Commander Suite — it validates whether "
    "a technically strong stock deserves your capital. It implements Minervini's 8-point growth "
    "scorecard, a Piotroski F-Score for corporate safety, and a 17-point Overall Fundamental Rating "
    "covering four domains. Apply on the candidate stock's daily chart after the screener identifies a setup."
)
note("RULE",
     "Positional trade: Overall Grade ≥ B (score ≥ 9/17). "
     "Swing trade: Overall Grade ≥ C (score ≥ 6/17). "
     "Grade D or F: only speculative, full-size exit at T1.", color=C_AMBER)

h2("22.2 Input Fields")
tbl(
    ["Setting","Default","Options","Description"],
    [
        ["Show Data Table","ON","ON/OFF","Toggle entire panel on/off"],
        ["Table Position","Bottom Right","9 positions","Where the panel renders on the chart"],
        ["Text Size","Small","Tiny/Small/Normal/Large/Auto","Readability preference"],
        ["Spacer Rows","0","0–50","Add blank rows below table for stacking with other indicators"],
        ["Background Color","Dark Navy","Color picker","Panel background"],
        ["Text Color","Light Blue-Gray","Color picker","All text in the panel"],
        ["Healthy Color","Green","Color picker","Metrics in healthy/passing range"],
        ["Warning Color","Amber","Color picker","Borderline metrics"],
        ["Danger Color","Red","Color picker","Failing/dangerous metrics"],
        ["Border Color","Dark Blue","Color picker","Panel border"],
    ],
    col_widths=[2.0,1.0,1.5,2.5]
)

h2("22.3 Section 1: Macro & Climate")
body("Live macroeconomic context — fetched from TradingView directly:")
tbl(
    ["Row","Metric","Green Signal","Red Signal","Trading Implication"],
    [
        ["India 10Y Yield","Current yield + trend","FALLING","RISING","Rising rates = equity PE compression. Avoid high-PE stocks."],
        ["USD/INR","Current rate + trend","STRENGTHENING (↓)","WEAKENING (↑)","Strengthening INR = FII inflows favorable to Indian markets"],
        ["CNX500 Trend","Broad market health","STRONG UPTREND","DOWNTREND","Context for all trades in current session"],
        ["Sector","Stock's sector name","Yellow informational","—","Confirms sector classification for RS analysis"],
    ],
    col_widths=[1.5,1.8,1.5,1.5,2.2]
)

h2("22.4 Section 2: Momentum — Revenue & Earnings")
tbl(
    ["Row","Metric","Green","Amber","Red","Minervini Insight"],
    [
        ["Market Cap","Total cap (INR format)","—","—","—","Reference only"],
        ["Revenue (TTM)","Annual turnover","—","—","—","Reference only"],
        ["Rev Growth (YoY)","Qtr rev vs year-ago",">20%",">0%","<0%","Must exceed 20% for top-quality positional"],
        ["NI Growth (YoY)","Qtr net income vs yr-ago",">25%",">0%","<0%","Earnings growth is the fuel for stock moves"],
        ["Net Income (TTM)","Trailing 12M profit",">0","—","<0","Must be positive — no exceptions for positionals"],
        ["EPS Diluted (FQ)","Latest quarter EPS",">0","—","<0","Earnings per share growing = institutional interest"],
        ["EPS Growth (YoY)","EPS vs year-ago qtr",">25%",">0%","<0%","Key Minervini metric — look for 3+ consecutive quarters"],
        ["Earnings Acceleration","Q growth > prior Q growth","YES","—","NO","MOST POWERFUL signal — compounding growth rate"],
    ],
    col_widths=[1.5,2.0,0.8,0.8,0.8,2.1]
)
note("KEY INSIGHT",
     "Earnings Acceleration (current Q growth > prior Q growth) combined with a technical breakout "
     "is the single most powerful setup in the stock market. If EPS Acceleration = YES, "
     "consider upsizing the position by 25–50% of normal allocation.", color=C_DKGRN)

h2("22.5 Section 3: Margins (TTM)")
tbl(
    ["Row","Metric","Green","Amber","Red"],
    [
        ["Gross Margin","(Gross Profit/Revenue)×100",">30%","15–30%","<15%"],
        ["EBITDA Margin","(EBITDA/Revenue)×100",">20%","10–20%","<10%"],
        ["Operating Margin","(Op Income/Revenue)×100",">15%","0–15%","<0%"],
        ["ROE (TTM)","Return on Equity (annualised)",">20%","12–20%","<12%"],
        ["ROA (TTM)","Return on Assets (annualised)",">10%","5–10%","<5%"],
        ["Accrual Ratio","Net Income vs Cash Flow","Negative (cash > NI)","—","Positive >10% (fake earnings)"],
    ],
    col_widths=[1.8,2.5,1.0,1.0,2.2]
)
note("RED FLAG",
     "If Accrual Ratio > 10%, the company is reporting Net Income but NOT generating cash. "
     "This is an earnings manipulation warning. Do not take positional trades on companies "
     "with Accrual Ratio > 10%.", color=C_RED)

h2("22.6 Section 4: Health & Value")
tbl(
    ["Row","Metric","Green","Amber","Red"],
    [
        ["Debt/Equity","Total Debt / Total Equity","<0.5×","0.5–1.5×",">1.5×"],
        ["Current Ratio","Current Assets / Current Liabilities",">2.0×","1.0–2.0×","<1.0×"],
        ["FCF (TTM)","Free Cash Flow = OCF − CapEx","Positive","—","Negative"],
        ["P/E Ratio","Price / TTM EPS","<25×","25–50×",">50× or negative"],
        ["P/B Ratio","Price / Book Value per Share","<1.5×","1.5–5.0×",">5.0×"],
    ],
    col_widths=[1.8,2.5,1.0,1.0,2.2]
)
note("DEBT TRAP",
     "If D/E > 1.5× AND India 10Y Yield is RISING — avoid. Rising rates crush leveraged companies. "
     "Heavily indebted companies have the most to lose in a rate-rise environment.", color=C_RED)

h2("22.7 Section 5: Minervini Fundamental Score (0–8)")
body("Eight binary pass/fail criteria targeting high-growth company characteristics:")
tbl(
    ["Criterion","Pass Threshold","Score","Why"],
    [
        ["Revenue Growth > 20%","YoY Rev Growth > 20%","+1","Institutional-quality revenue expansion"],
        ["NI Growth > 25%","YoY Net Income Growth > 25%","+1","Profitability accelerating faster than revenue"],
        ["Accelerating EPS","Current Q growth > prior Q growth","+1","Compounding earnings momentum — the most powerful signal"],
        ["ROE > 15%","ROE (TTM) > 15%","+1","Capital efficiency — management creating value"],
        ["Gross Margin > 15%","Gross Margin (TTM) > 15%","+1","Pricing power and business quality floor"],
        ["D/E < 1.5","Debt/Equity < 1.5×","+1","Balance sheet durability through market cycles"],
        ["Current Ratio > 1.0","Current Assets / Current Liabilities > 1.0","+1","Short-term solvency — can survive a correction"],
        ["FCF Positive","Free Cash Flow (TTM) > 0","+1","Self-funding business — not dependent on markets for survival"],
    ],
    col_widths=[2.0,2.2,0.6,3.2]
)
tbl(
    ["Score","Stars","Quality","Positional Trade?","Swing Trade?"],
    [
        ["7–8","★★★★★ (5 Stars)","Institutional-grade growth engine","Full size — this is the sweet spot","Yes"],
        ["5–6","★★★★☆ (4 Stars)","Strong fundamentals","Full size","Yes"],
        ["3–4","★★★☆☆ (3 Stars)","Adequate — momentum driven","Half size or wait for earnings improvement","Yes"],
        ["1–2","★★☆☆☆ (2 Stars)","Weak fundamentals","Avoid positional — short swing only","With tight stop"],
        ["0","★☆☆☆☆ (1 Star)","Fundamentally broken","Never","Intraday/scalp only"],
    ],
    col_widths=[0.6,1.5,2.0,2.5,1.0]
)

h2("22.8 Section 6: Piotroski F-Score (0–9)")
body(
    "The Piotroski F-Score measures corporate safety across nine binary checks across "
    "profitability, leverage, and operating efficiency. It identifies companies at risk of "
    "financial distress BEFORE it shows up in the stock price."
)
tbl(
    ["Score","Stars","Interpretation","Action"],
    [
        ["8–9","★★★★★","Corporate fortress — resilient, improving, safe","Hold through corrections. Widest trailing stop."],
        ["6–7","★★★★☆","Solid business — financially stable","Standard hold. Normal trailing stops."],
        ["4–5","★★★☆☆","Average — minor concerns","Monitor debt and FCF quarterly. Tighter stops."],
        ["2–3","★★☆☆☆","Potential distress signals","Swing only. Hard stops. Never add on pullback."],
        ["0–1","★☆☆☆☆","High distress risk","Do NOT hold. Scalp only or avoid completely."],
    ],
    col_widths=[0.6,1.2,2.8,3.4]
)

h2("22.9 Section 7: Overall Fundamental Rating (0–17)")
body("A holistic composite scoring four domains with stricter thresholds than the 8-pt Minervini score:")
tbl(
    ["Domain","Max Points","Criteria (1 point each unless noted)"],
    [
        ["Macro","3","CNX500 Strong Uptrend (+1) · India 10Y Yield Falling (+1) · INR Strengthening (+1)"],
        ["Momentum","5","Rev YoY>20% · NI YoY>25% · EPS YoY>25% · EPS Accelerating · Net Income Positive"],
        ["Margins","4","Gross Margin>30% · EBITDA>20% · Operating Margin>15% · ROE>20%"],
        ["Health & Value","5","D/E<1.0 · Current Ratio>1.5 · FCF Positive · P/E<40 · P/B<5.0"],
    ],
    col_widths=[1.5,1.0,5.5]
)
tbl(
    ["Score","Grade","Meaning","Trade Horizon Recommendation"],
    [
        ["15–17","A+ EXCEPTIONAL","Institutional-grade quality — rare","Full positional. Hold 6–18 months. Loose trailing stop (50-DMA)."],
        ["12–14","A STRONG","Excellent fundamentals","Full positional. Hold 3–12 months."],
        ["9–11","B GOOD","Solid, tradeable quality","Positional or swing. Standard management."],
        ["6–8","C FAIR","Acceptable momentum play","Swing only. Tight trailing stops."],
        ["3–5","D WEAK","Avoid for positional trades","Scalp/momentum only. 50% position. Hard stop at T1."],
        ["0–2","F POOR","Financially dangerous","Never trade for position. Avoid completely."],
    ],
    col_widths=[0.8,1.5,2.0,3.7]
)

h2("22.10 The 15-Second X-Ray Checklist")
body("When a technical trigger flashes, run this rapid fundamental validation:")
bullet("1. What is the Overall Grade? ≥ B = proceed. D/F = halve size or skip.")
bullet("2. Is it safe? Piotroski > 4, FCF positive, D/E < 1.5×?")
bullet("3. Is it a hyper-grower? EPS Accelerating = YES AND Minervini Score ≥ 6?")
bullet("4. What is the play horizon? A/A+ = Positional (marry it). B/C = Swing (date it). D/F = Scalp (one-night stand).")

divider()

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 23 — SWING ZIGZAG STRICT v6.0
# ═══════════════════════════════════════════════════════════════════════════
h1("Section 23 — Weinstein Swing Zigzag Strict v6.0: Complete Reference")

h2("23.1 Purpose")
body(
    "The Swing Zigzag visualises pure market structure — Higher Highs (HH), Higher Lows (HL), "
    "Lower Lows (LL), Lower Highs (LH) — using confirmed pivot points. It answers the most "
    "fundamental question in trading: Is this stock in a valid uptrend structure before I enter? "
    "It also provides Fibonacci retracement levels, a structural trend state HUD, and Break of "
    "Structure (BoS) alerts for hands-free monitoring."
)
body("Apply on Daily (swing trades) or Weekly (positional trades) chart alongside Dashboard and Strategy.")

h2("23.2 All Input Fields")
h3("MTF Pivot Settings")
tbl(
    ["Input","Default","Auto-Selected For","Notes"],
    [
        ["Monthly Pivot Length","1","Monthly charts","Bars left+right for monthly pivot detection"],
        ["Weekly Pivot Length","5","Weekly charts","5 bars left + 5 right = confirmed weekly pivots. Key setting."],
        ["Daily Pivot Length","2","Daily charts","2 bars left + 2 right for daily pivots. Balanced sensitivity."],
        ["Intraday Pivot Length","2","<Daily charts","For intraday analysis if needed"],
    ],
    col_widths=[2.2,0.8,1.8,3.2]
)
body("The indicator auto-selects the correct pivot length based on chart timeframe. No manual adjustment needed.")

h3("Display Settings")
tbl(
    ["Input","Default","Description"],
    [
        ["Show Trend State Panel","ON","Displays the HUD info panel (TREND, STRUCTURE, SWING COUNT, etc.)"],
        ["Show Projection Line","ON","Draws dashed projection line from current live pivot — unconfirmed swing direction"],
        ["Panel Position","Top Right","Where the HUD panel renders on chart"],
        ["Choppiness Lookback (weeks)","52","Rolling 52-week window for counting direction flips"],
    ],
    col_widths=[2.5,0.8,4.7]
)

h3("Fibonacci Levels")
tbl(
    ["Input","Default","Description"],
    [
        ["Show 50% & 61.8% Retracement","ON","Draws Fibonacci levels for the current confirmed swing"],
        ["Extend Lines (bars right)","20","How far right the Fib lines extend past the current bar"],
    ],
    col_widths=[2.5,0.8,4.7]
)

h3("Major / Minor Pivot Settings")
tbl(
    ["Input","Default","Description"],
    [
        ["Major Pivot Mode","Auto","Auto detects: ETF = ≥4% swing threshold, Stock = ≥8% swing threshold. Custom = manual."],
        ["Custom Threshold (%)","8.0%","Used only when Major Pivot Mode = 'Custom'"],
    ],
    col_widths=[2.5,0.8,4.7]
)

h2("23.3 Reading the Chart Visuals")
tbl(
    ["Visual","Meaning","Trading Implication"],
    [
        ["Green zigzag lines","HH after HL (confirmed uptrend leg)","Trend is healthy — hold longs, buy pullbacks"],
        ["Red zigzag lines","LH after LL (confirmed downtrend leg)","Avoid longs — structural breakdown confirmed"],
        ["White/faded lines","Mixed structure (broadening or sideways)","No clear trend — wait for resolution"],
        ["Dotted gray horizontal","Lock line — confirmed pivot level","The structural reference level for BoS"],
        ["Dashed projection line","Current UNCONFIRMED swing direction","Directional bias — not yet committed"],
        ["HH (+12.3%)","Higher High, X% gain from prior low","Uptrend confirmed leg — percentage is the move size"],
        ["HL (−5.1%)","Higher Low, X% pullback from prior high","Healthy pullback depth — under 15% = ideal"],
        ["EH / EL","Equal High / Equal Low (within 0.1% tolerance)","Double top/bottom — SIDEWAYS warning. Stop buying breakouts."],
        ["THICK lines (width 3)","Major pivot — swing ≥ 8% (stock) or ≥ 4% (ETF)","Institutional-level structural move. Full size entry."],
        ["Thin lines (width 1)","Minor pivot — below threshold","Minor noise. Half-size entry maximum."],
        ["Solid Fib lines (purple/orange)","Trend is UP, BoS confirmed — levels are ACTIONABLE","Enter at Fib levels with reversal candle confirmation"],
        ["Dotted Fib lines","Trend is DOWN or structure unconfirmed","Reference only — DO NOT enter at dotted Fib levels"],
    ],
    col_widths=[2.2,2.5,2.8]
)

h2("23.4 The HUD Panel — Field-by-Field Guide")
tbl(
    ["Field","Values","How to Interpret"],
    [
        ["TREND","UPTREND / DOWNTREND / SIDEWAYS","Must show UPTREND for all bull market entries. SIDEWAYS = wait."],
        ["STRUCTURE","HH/HL | LH/LL | HL/HH | etc.","HH/HL = healthy uptrend. LH/LL = downtrend. Mixed = no clear bias."],
        ["SWING COUNT","1 swing, 2 swings, 3 swings, etc.","1–2 swings = fresh trend, ideal for breakout entries. 3+ = mature trend, prefer pullbacks. 4–5 = exhaustion risk, reduce/exit."],
        ["AMPLITUDE","e.g., '23.5%'","% change from locked low to locked high. Larger = more powerful trend leg."],
        ["CHOPPINESS","e.g., '3 flips / 52W'","≤4 flips = GREEN (trending). 5–8 = YELLOW (caution). >8 = RED (abandon — stop-hunt environment)."],
        ["CONFIG","e.g., 'Pivot: L=5 R=5'","Active pivot length (auto-selected by chart timeframe)"],
    ],
    col_widths=[1.8,2.0,4.2]
)
note("CRITICAL RULE",
     "If CHOPPINESS > 8 (red), do not enter ANY setup on this chart. "
     "High choppiness means the stock routinely hunts stops and reverses. "
     "Your carefully placed stop loss will be triggered before the move develops.", color=C_RED)

h2("23.5 Entry Setup 1 — Structural Breakout (BoS)")
body("Aggressive entry type — hunting momentum as it ignites:")
bullet("1. Ensure the HUD shows lockedHigh (the structural resistance level to break)")
bullet("2. Watch the dashed Projection Line — it crosses the lockedHigh value when a Bullish BoS fires")
bullet("3. Entry: At BoS confirmation (volume surge) or on close above lockedHigh for safer entry")
bullet("4. Stop: Just below the most recent lockedLow — do not give excessive room below this")
bullet("5. Enable the 'Bullish BoS' TradingView alert to receive phone notifications")
note("TIP", "If volume is immense at the BoS moment, you may enter before the daily close. "
     "For lower-conviction setups, always wait for daily close above the pivot.", color=C_AMBER)

h2("23.6 Entry Setup 2 — Golden Pocket Pullback (Fibonacci)")
body("High-probability, lower-risk setup using the automated Fibonacci system:")
bullet("1. HUD must show UPTREND and Fib lines must be SOLID (not dotted)")
bullet("2. Wait for price to retrace from the recent HH back into the Fib zone:")
bullet("   Purple line = 50% retracement (standard entry zone)", 1)
bullet("   Orange line = 61.8% retracement (Golden Pocket — best entry zone)", 1)
bullet("3. Do NOT blindly place a limit at the Fib line — wait for a reversal candle (Hammer or Bullish Engulfing) confirming rejection at the level")
bullet("4. Stop: Below the lockedLow (the swing anchor). If this is breached, bullish thesis is mathematically invalid.")
note("RULE", "Dotted Fib lines = trend is down or unconfirmed. Never enter at dotted Fib levels. "
     "This is the most common Fibonacci mistake — trading broken structures.", color=C_RED)

h2("23.7 Active Trade Management With Zigzag")
bullet("1. Riding the Trend: As price rises, the indicator prints new HH and HL labels")
bullet("2. Trailing Stop: Every time a new HL (Higher Low) is formally locked in, move your broker stop to 1 tick below that new HL level")
bullet("3. Swing Count 4–5: SELL 50–75% of position regardless of trailing stop — structural exhaustion probable")
bullet("4. Bearish BoS: If live price drops below the last lockedLow → exit full position immediately. Structure is broken.")

h2("23.8 Pivot Classification Reference")
tbl(
    ["This Pivot","After This Pivot","Signal","Confidence"],
    [
        ["HH (Higher High)","HL (Higher Low) — confirmed","Uptrend continuation","High — accumulate on pullback"],
        ["HL (Higher Low)","HH (Higher High) — confirmed","Uptrend continuation","High — hold existing longs"],
        ["LH (Lower High)","LL (Lower Low) — confirmed","Downtrend continuation","High — exit longs immediately"],
        ["LL (Lower Low)","LH (Lower High) — confirmed","Downtrend continuation","High — avoid all longs"],
        ["HH after LL","—","Broadening structure","Low — avoid entries; structure invalid"],
        ["EH (Equal High)","—","Sideways / double top","Medium — trend downgraded to SIDEWAYS"],
        ["EL (Equal Low)","—","Sideways / double bottom","Medium — wait for resolution"],
    ],
    col_widths=[1.8,2.2,2.0,2.0]
)

divider()

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 24 — BETA SCREENER v2.3 COMPLETE REFERENCE
# ═══════════════════════════════════════════════════════════════════════════
h1("Section 24 — Commander Screener Beta Edition v2.3: Complete Reference")

h2("24.1 Purpose & Architecture")
body(
    "The Beta Screener v2.3 is a 'headless' indicator — it outputs 26 precisely calibrated data "
    "columns to TradingView's Data Window and Stock Screener tool, rather than drawing a visible chart. "
    "This makes it compatible with TradingView's Screener, allowing you to filter thousands of stocks "
    "by Stage, Catalyst, Alpha Score, and 23 other metrics simultaneously."
)
body("Single-stock deep-dive on its chart. Use alongside Beta Edition's Screener output in TradingView Screener tool.")

h2("24.2 Settings — Keep Synced With Dashboard v63.3")
body(
    "All Beta Edition settings must match the Dashboard v63.3 values for signal consistency. "
    "If these diverge, you will see different Stage/RS/Catalyst readings between the two tools."
)
tbl(
    ["Setting","Default","Sync With"],
    [
        ["Weinstein SMA Length (Weekly)","30","Dashboard v63.3 → 30W MA Length"],
        ["30 WMA Slope Lookback (Weeks)","4","Dashboard v63.3 → Slope Lookback"],
        ["30WMA Slope Threshold","0.0005","Dashboard v63.3 → Slope Threshold"],
        ["Benchmark (Nifty 50)","NSE:NIFTY","Dashboard v63.3 → Benchmark 1"],
        ["Benchmark (Nifty 500)","NSE:CNX500","Dashboard v63.3 → Benchmark 2"],
        ["Mansfield RS Length","26 weeks","Dashboard v63.3 → Mansfield RS Length"],
        ["RS Slope Lookback","8 weeks","Dashboard v63.3 → RS Slope Lookback"],
        ["50 DMA Length","50","Dashboard v63.3 → 50 DMA Length"],
        ["150 DMA Length","150","Dashboard v63.3 → 150 DMA Length"],
        ["200 DMA Length","200","Dashboard v63.3 → 200 DMA Length"],
        ["Macro Edge: Institutional Vol Bias","ON","Strategy v4.5 → Macro Edge"],
    ],
    col_widths=[2.8,1.2,3.0]
)

h2("24.3 All 26 Data Window Outputs")
tbl(
    ["#","Column Name","Range/Values","How to Use in Screener"],
    [
        ["1","Alpha Score","0–100+","Primary filter: ≥ 4 stars (80+) for STRONG BUY"],
        ["2","Dashboard Quality Score","0–100","Composite score — sort descending for best setups"],
        ["3","Alpha Stars (0–5)","0, 1, 2, 3, 4, 5","Set ≥ 4 to filter top-quality only"],
        ["4","Confluence (0–6 Stars)","0–6","Price/Vol/EMA/RSI/ADX/Structure. ≥ 4 = high confluence"],
        ["5","Stage (Numeric)","1.0, 2.0, 2.1, 3.0, 4.0","Set 1.0 ≤ Stage ≤ 2.1 (Stage 1 base to Stage 2 up)"],
        ["6","Persona (0–5)","0=Lag, 2=Vol, 3=Turn, 4=Mom, 5=Lead","Filter: Persona ≥ 4 for Leaders and Momentum stocks"],
        ["7","Style (0–3)","0=Wait, 1=Pos, 2=Swing, 3=Both","Filter: Style = 3 (Both) for maximum conviction setups"],
        ["8","Catalyst ID (0–6)","0=None,1=ACCUM,2=POS-BO,3=PB,4=VCP,5=REV,6=GAP","Filter: > 0 for any catalyst. Set = 2 for positional breakouts only."],
        ["9","Recommended SL Price","Price","Use as hard stop loss — do not recalculate manually"],
        ["10","Distance to SL (%)","% distance","< 5%: normal size. > 8%: halve position size."],
        ["11","Target 1 (2.0R)","Price","First profit target"],
        ["12","Target 2 (3.0R)","Price","Runner profit target"],
        ["13","Warn: T1 Hits Resistance","0 or 1","MUST be 0 before entry. 1 = T1 is blocked, skip trade."],
        ["14","Signal: True Breakout","0 or 1","POS-BO breakout signal active today"],
        ["15","Signal: Pullback Hit","0 or 1","SWG-PB pullback signal active today"],
        ["16","Minervini Trend Template","0 or 1","1 = full 6-criteria Minervini template passing"],
        ["17","RS Value (vs Nifty 50)","Float (Mansfield)","Positive = outperforming Nifty 50"],
        ["18","RS Value (vs Nifty 500)","Float (Mansfield)","Primary RS check. Must be > 0 for entries."],
        ["19","Sector RS Value","Float (Mansfield)","Positive = outperforming own sector"],
        ["20","Vol Shelf (VWMA50>SMA50)","0 or 1","1 = institutional volume participation confirmed"],
        ["21","Daily Trend","1 or −1","1 = bullish daily trend"],
        ["22","Weekly Trend","1 or −1","1 = bullish weekly trend"],
        ["23","Current Price","Price","Reference"],
        ["24","Distance to 20 EMA (%)","% (±)","−2% to +2% = in pullback zone for SWG-PB"],
        ["25","Relative Volume (x-avg)","Multiplier","≥ 1.5 = institutional volume surge"],
        ["26","Anti-Algo SWG-BO Gate","0 or 1","MUST be 1 for breakout entries. 0 = HFT trap, skip."],
    ],
    col_widths=[0.4,2.3,1.5,3.8]
)

h2("24.4 Recommended Screener Filter Combinations")
h3("Filter 1: Strong Buy Positional Scan")
code("Stage = 2.1 (UP)  AND  Catalyst ID = 2 (POS-BO)  AND  Alpha Stars >= 4")
code("AND  Minervini Template = 1  AND  T1 Resistance Warning = 0  AND  Vol Shelf = 1")
body("Result: Only full-alignment Stage 2 breakouts with institutional volume and clear room to T1.")

h3("Filter 2: Swing Pullback Scan")
code("Signal: Pullback Hit = 1  AND  Weekly Trend = 1  AND  Stage >= 2.0")
code("AND  Vol Shelf = 1  AND  Distance to 20 EMA < 3%  AND  T1 Resistance Warning = 0")
body("Result: EMA20 pullback candidates in Stage 2 with institutional backing.")

h3("Filter 3: VCP Breakout Scan")
code("Catalyst ID = 4 (SWG-BO)  AND  Minervini Template = 1  AND  Anti-Algo Gate = 1")
code("AND  Alpha Stars >= 4  AND  Vol Shelf = 1  AND  RS (vs N500) > 0")
body("Result: High-quality VCP breakouts that passed the anti-HFT filter.")

h3("Filter 4: Full Market Scan (Any Quality Entry Today)")
code("Catalyst ID > 0  AND  Stage >= 1.0 AND Stage <= 2.1  AND  Alpha Stars >= 3")
code("AND  T1 Resistance Warning = 0  AND  Anti-Algo Gate = 1")
body("Result: Any catalyst firing today across Stage 1 and Stage 2 stocks with clear risk/reward.")

h2("24.5 EOD Scanning Workflow")
bullet("Step 1: Run screener 30 minutes before market close or immediately after close")
bullet("Step 2: Apply Filter 1 for positional setups, Filter 2 for swing pullbacks")
bullet("Step 3: Check Anti-Algo Gate = 1 on all breakout candidates")
bullet("Step 4: Verify Distance to SL. If > 8%: halve position size before proceeding")
bullet("Step 5: Pull top 3–5 candidates into Dashboard v63.3 for full analysis")

divider()

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 25 — COMMANDER WEB v3.0
# ═══════════════════════════════════════════════════════════════════════════
h1("Section 25 — Weinstein Commander Web v3.0: Complete Reference")

h2("25.1 Architecture & Purpose")
body(
    "Commander Web v3.0 is a Streamlit-based desktop trading command centre that integrates "
    "every component of the ecosystem: live Dhan broker connectivity, real-time market health, "
    "AI-powered grading, risk analytics, sector rotation analysis, options data, and portfolio "
    "management — all in a single browser-based interface."
)
body("Launch: python -m streamlit run weinstein_commander_web_v3.0.py")

h2("25.2 Navigation Pages")
tbl(
    ["Page","Purpose","Key Actions"],
    [
        ["DASHBOARD","Portfolio command centre — live P&L, market health, noise risk, deployed capital","Daily monitoring, market health check, noise risk review"],
        ["HUNTER","Trade discovery: Chartink scanners, enrichment with AI grades, candidate selection","Weekly EOD scan workflow, shortlist generation"],
        ["WATCHLIST","Watchlist generation and cloud sync","Build and maintain weekly watchlist"],
        ["COMMAND","Active Operations — open position management and trade ledger","Position monitoring, SL updates, scale-in/out decisions"],
        ["AI LAB","Pre-Flight checker, generative analysis, automated workflows","Trade validation, AI grading, batch analysis"],
        ["MACRO","Macro indicators: India VIX, USD/INR, Crude, Gold, US 10Y, Nifty 50","Context check before entering trades"],
        ["OPTIONS","NSE option chain data for selected symbols","Gauging OI, PCR, max pain for FII positioning"],
        ["AUTOPSY","Closed trade performance analytics — Sharpe, Drawdown, Win Rate, Expectancy","Monthly performance review, system optimisation"],
        ["BACKTEST","Strategy backtesting interface","Validate rule changes before live trading"],
        ["JOURNAL","Opens dhan_journal_v7.py — full trade journal (Streamlit app)","Logging trades, reviewing historical decisions"],
        ["X-RAY","Opens fundamental_xray.py — Fundamental X-Ray standalone app","Deep fundamental analysis on demand"],
        ["TV SIDECAR","Opens tv_sidecar_app.py — TradingView integration","Chart synchronisation, Pine code management"],
    ],
    col_widths=[1.5,3.2,3.3]
)

h2("25.3 DASHBOARD Page — Status Bar")
body("The top status bar shows five key metrics at all times:")
tbl(
    ["Cell","Content","Green = Good","Red = Bad"],
    [
        ["System","SYSTEM ONLINE / AUTH EXPIRED / API OFFLINE","SYSTEM ONLINE","AUTH EXPIRED = re-authenticate Dhan"],
        ["Market","BULLISH / BEARISH (<200DMA)","BULLISH","BEARISH = no new longs"],
        ["Noise Risk","✔ SECURE / ⚠ N AT RISK","✔ SECURE","⚠ N = N positions have SL < 1.5×ATR"],
        ["Deployed Capital","Amount deployed + %","< 70% deployed","> 80% = over-committed"],
        ["Available Balance","Cash in Dhan account","Adequate for new trades","Low = no capacity for new entries"],
    ],
    col_widths=[1.5,2.5,2.0,2.0]
)

h2("25.4 HUNTER Page — Trade Discovery Workflow")
h3("Sub-tab: SCANNERS")
body("Runs three recovery scanners and loads Chartink results. Equivalent to chartink_scanner_pro.py output.")
bullet("Runs Scanner 5 (REV-RS), Scanner 6 (REV-CB), Scanner 7 (REV-EARLY) for recovery market")
bullet("Also supports bull market screeners integrated with Commander Beta Edition output")
bullet("Results displayed in ranked table with Symbol, Stage, RS, Catalyst columns")

h3("Sub-tab: ENRICHMENT")
body("Takes the scanner output and enriches each candidate with AI grading:")
bullet("Calls get_weinstein_score() from ai_grading_engine.py for each symbol")
bullet("Shows: Star Rating, Quant Grade (A–F), Quant Score (/100), Breakdown by component, Catalyst tag")
bullet("Cached results load instantly for previously graded symbols")
bullet("Parallel threading: ThreadPoolExecutor fetches multiple symbols simultaneously")

h3("Sub-tab: SELECTION")
body("Final shortlist view combining technical + fundamental + AI grades:")
bullet("Filter by minimum Alpha Score, Catalyst type, and Stage")
bullet("One-click add to Watchlist for selected candidates")

h2("25.5 COMMAND Page — Active Position Management")
h3("Sub-tab: ACTIVE OPS")
body("Real-time monitoring of all open positions with live LTP from Dhan holdings:")
tbl(
    ["Column","Description"],
    [
        ["Symbol","Stock ticker"],
        ["Entry/LTP","Your buy price vs current live price"],
        ["P&L (₹ and %)","Unrealised profit/loss"],
        ["Current R","Profit in R-multiples: (LTP−Entry)/(Entry−SL)"],
        ["SL / ATR","Your stop loss and current ATR14"],
        ["Noise Flag","⚠ = SL < 1.5×ATR (invisible stop risk)"],
        ["Days Held","Calendar days since entry"],
        ["AI Grade","Star rating from grading engine"],
        ["Action","HOLD / TRIM / SCALE IN / STOP HIT — driven by Dashboard logic"],
    ],
    col_widths=[2.0,5.0]
)

h3("Sub-tab: LEDGER")
body("Closed trade history with entry/exit details, P&L, and quality assessment.")

h2("25.6 AI LAB Page")
h3("Sub-tab: PRE-FLIGHT")
body("Comprehensive trade validation checklist before entry. Runs 10-gate check:")
bullet("Gate 1: Market health (not BEARISH)")
bullet("Gate 2: Stage verification (Stage 2 preferred)")
bullet("Gate 3: RS check (Mansfield RS > 0 vs CNX500)")
bullet("Gate 4: AI Grade (≥ B for positional)")
bullet("Gate 5: Noise risk (SL ≥ 1.5×ATR14)")
bullet("Gate 6: Portfolio capacity (slots available)")
bullet("Gate 7: Correlation check (not shadow pair with existing position)")
bullet("Gate 8: Regime gate (bull/correction/bear)")
bullet("Gate 9: Earnings proximity (not within 3 days)")
bullet("Gate 10: T1 clear room (not blocked by overhead resistance)")
body("All 10 gates must show GREEN before entry. Any RED = do not enter.")

h3("Sub-tab: MACRO")
body("Macro overview with India VIX, USD/INR, Crude Oil, Gold, US 10Y yield, and Nifty 50:")
bullet("Each shows: LTP, SMA50, SMA200, 1M%, 1Y%, Percentile rank, Stage")
bullet("VIX > 20: heightened fear — consider reducing new entries")
bullet("India 10Y Yield rising: PE compression risk — avoid high-P/E stocks")
bullet("USD/INR weakening: FII outflow risk — reduce exposure to FII-heavy large-caps")

h2("25.7 AUTOPSY Page — Performance Analytics")
body(
    "The AUTOPSY page runs compute_portfolio_analytics() on the closed trade database "
    "and displays comprehensive performance metrics:"
)
tbl(
    ["Metric","Definition","Target"],
    [
        ["Total Trades","Count of closed trades","≥ 20 for statistical significance"],
        ["Win Rate","% of profitable trades","≥ 45%"],
        ["Avg Win / Avg Loss","Average winner vs loser in rupees","Ratio ≥ 2.0×"],
        ["Profit Factor","Sum of wins / Sum of losses","≥ 1.5 (> 2.0 = excellent)"],
        ["Expectancy","Expected profit per trade (rupees)","Positive — the more the better"],
        ["Max Drawdown","Peak-to-trough equity loss (₹ and %)","< 15% of capital"],
        ["Sharpe Ratio","Risk-adjusted return vs India T-bill (6.5%)","≥ 1.5"],
        ["Sortino Ratio","Downside-adjusted return","≥ 2.0"],
        ["Total Realized P&L","Net profit from all closed trades","Track vs CNX500 benchmark"],
    ],
    col_widths=[2.0,3.0,2.5]
)
note("MONTHLY REVIEW",
     "Run AUTOPSY once per month. If Expectancy is negative over 20+ trades, "
     "pause trading and audit your entry criteria. If Win Rate < 35%, review "
     "your stop placement — stops may be too tight.", color=C_AMBER)

h2("25.8 Sector Rotation Analysis (fetch_sector_momentum_cached)")
body(
    "The MACRO page includes a sector rotation panel covering 11 NSE sector indices. "
    "It calculates RS momentum, acceleration, and RRG quadrant for each sector:"
)
tbl(
    ["Column","Description"],
    [
        ["Sector","Nifty sector name"],
        ["4W %","4-week price return"],
        ["Prior 4W %","Prior 4-week return (for acceleration comparison)"],
        ["8W %","8-week return (RS measure)"],
        ["Acceleration","Recent 4W − Prior 4W (FORM-01 fix)"],
        ["RS-Ratio","Sector RS vs CNX500 benchmark (RRG X-axis)"],
        ["RRG Quadrant","🟢 Leading / 🟡 Weakening / 🔵 Improving / 🔴 Lagging"],
        ["Signal","🟢 Accelerating (accel>2) / 🔴 Decelerating (accel<−2) / 🟡 Neutral"],
    ],
    col_widths=[2.0,5.0]
)
body("Sorted by Acceleration descending — the top sectors are gaining momentum right now.")
body("Use this to prioritise which sectors to scan first in the weekly Hunter workflow:")
bullet("🟢 Leading + Accelerating = highest priority — stocks here have the wind at their backs")
bullet("🔵 Improving = second priority — early rotation candidates, hidden accumulation")
bullet("🟡 Weakening = watch for exits — leadership fading")
bullet("🔴 Lagging = avoid new entries regardless of individual stock setup quality")

doc.add_page_break()
doc.save(OUTPUT)
print(f"Part 6 saved: {OUTPUT}")
print("Sections 22–25 complete.")
