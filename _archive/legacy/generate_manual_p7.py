"""
Recovery Trading Manual — Generator Part 7
BULL MARKET SECTION: Sections 26–30
- Section 26: The Complete Bull Market Workflow (End-to-End)
- Section 27: Bull Market Risk Management Framework
- Section 28: Common Mistakes & How to Avoid Them
- Section 29: Psychology — The Trader's Mindset
- Section 30: Bull Market Quick Reference Card
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = r"C:\Users\jayra\Documents\GeminiVSCode\Recovery_Trading_Manual.docx"
doc = Document(OUTPUT)

C_NAVY   = RGBColor(0x1F, 0x35, 0x64)
C_TEAL   = RGBColor(0x00, 0x70, 0x70)
C_AMBER  = RGBColor(0xC5, 0x5A, 0x11)
C_GREEN  = RGBColor(0x37, 0x86, 0x30)
C_LGREY  = RGBColor(0xF2, 0xF2, 0xF2)
C_DGREY  = RGBColor(0x40, 0x40, 0x40)
C_RED    = RGBColor(0xC0, 0x00, 0x00)
C_PURPLE = RGBColor(0x5B, 0x2C, 0x6F)
C_DKGRN  = RGBColor(0x14, 0x5A, 0x32)

def sf(run,bold=False,italic=False,size=11,color=C_DGREY):
    run.bold=bold;run.italic=italic;run.font.size=Pt(size);run.font.color.rgb=color

def h1(text):
    p=doc.add_paragraph();p.paragraph_format.space_before=Pt(18);p.paragraph_format.space_after=Pt(6)
    r=p.add_run(text);sf(r,bold=True,size=18,color=C_NAVY)

def h2(text):
    p=doc.add_paragraph();p.paragraph_format.space_before=Pt(12);p.paragraph_format.space_after=Pt(4)
    r=p.add_run(text);sf(r,bold=True,size=14,color=C_TEAL)

def h3(text):
    p=doc.add_paragraph();p.paragraph_format.space_before=Pt(8);p.paragraph_format.space_after=Pt(2)
    r=p.add_run(text);sf(r,bold=True,size=12,color=C_NAVY)

def body(text,indent=0):
    p=doc.add_paragraph();p.paragraph_format.space_after=Pt(4)
    p.paragraph_format.left_indent=Inches(indent*0.25)
    r=p.add_run(text);sf(r,size=10.5)

def bullet(text,level=0):
    p=doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent=Inches(0.25+level*0.2);p.paragraph_format.space_after=Pt(2)
    r=p.add_run(text);sf(r,size=10.5)

def note(label,text,color=C_AMBER):
    p=doc.add_paragraph()
    r1=p.add_run(f"{label}: ");sf(r1,bold=True,size=10.5,color=color)
    r2=p.add_run(text);sf(r2,size=10.5)

def code(text):
    p=doc.add_paragraph();p.paragraph_format.left_indent=Inches(0.3)
    r=p.add_run(text);sf(r,bold=True,size=9.5,color=C_NAVY)

def divider():
    p=doc.add_paragraph();p.paragraph_format.space_before=Pt(6);p.paragraph_format.space_after=Pt(6)
    r=p.add_run("─"*85);sf(r,size=8,color=RGBColor(0xCC,0xCC,0xCC))

def shade_cell(cell,rgb):
    tc=cell._tc;tcPr=tc.get_or_add_tcPr()
    shd=OxmlElement('w:shd')
    shd.set(qn('w:fill'),f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}")
    shd.set(qn('w:val'),'clear');tcPr.append(shd)

def tbl(headers,rows,col_widths=None):
    t=doc.add_table(rows=1,cols=len(headers))
    t.style='Table Grid';t.alignment=WD_TABLE_ALIGNMENT.CENTER
    hdr=t.rows[0]
    for i,h in enumerate(headers):
        cell=hdr.cells[i];shade_cell(cell,C_NAVY)
        p=cell.paragraphs[0];r=p.add_run(h);sf(r,bold=True,size=10,color=RGBColor(0xFF,0xFF,0xFF))
    for ri,rd in enumerate(rows):
        row=t.add_row()
        bg=C_LGREY if ri%2==0 else RGBColor(0xFF,0xFF,0xFF)
        for ci,val in enumerate(rd):
            cell=row.cells[ci];shade_cell(cell,bg)
            p=cell.paragraphs[0];r=p.add_run(str(val));sf(r,size=10)
    if col_widths:
        for i,w in enumerate(col_widths):
            for row in t.rows:row.cells[i].width=Inches(w)
    doc.add_paragraph()
    return t

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 26 — THE COMPLETE BULL MARKET WORKFLOW
# ═══════════════════════════════════════════════════════════════════════════
h1("Section 26 — The Complete Bull Market Workflow: End-to-End")

h2("26.1 Overview: Five-Phase Process")
body(
    "Every bull market trade goes through a five-phase pipeline. "
    "Skipping any phase increases the probability of entering a bad trade. "
    "The system is designed so each phase eliminates most candidates, leaving only the best."
)
tbl(
    ["Phase","Tool","Goal","Time Required"],
    [
        ["1. EOD Screening","Beta Screener v2.3 / Commander Web HUNTER","Filter 2,000+ stocks to 10 quality candidates","20–30 min (Sunday)"],
        ["2. Fundamental Gate","Fundamental X-Ray v2.2","Validate business quality — eliminate junk","15–20 min"],
        ["3. Structure Validation","Swing Zigzag v6.0 + Dashboard v63.3","Confirm uptrend structure and catalyst","10–15 min"],
        ["4. Execution","Strategy v4.5 Risk Allocator + Pre-Flight","Precise entry/stop/size and 10-gate check","5–10 min per trade"],
        ["5. Active Management","Dashboard v63.3 + Commander Web COMMAND","Monitor, trail, scale, and exit systematically","5–10 min daily"],
    ],
    col_widths=[1.2,2.3,2.8,1.7]
)

h2("26.2 PHASE 1: Weekly EOD Screening (Every Sunday, 30 minutes)")

h3("Step 1: Market Health Check (5 min)")
body("Open Commander Web → DASHBOARD. Check the status bar:")
bullet("Market cell shows BULLISH → proceed with full workflow")
bullet("Market cell shows BEARISH → no new longs. Review exits on open positions only.")
bullet("Correction (CNX500 below SMA50 but above SMA200) → reduce size 25%, swing trades only")
body("Also check Commander Web MACRO page:")
bullet("India VIX > 20 → heightened volatility, reduce all position sizes by 25%")
bullet("USD/INR weakening trend → FII outflow risk, prefer domestic-focused sectors")
bullet("Sector RRG table → identify top 2–3 'Leading + Accelerating' sectors for this week")

h3("Step 2: Sector Prioritisation (5 min)")
body("From the Commander Web MACRO → Sector Rotation table, identify:")
bullet("Column 'RRG Quadrant' = 🟢 Leading AND Column 'Signal' = 🟢 Accelerating → Priority A sectors")
bullet("Column 'RRG Quadrant' = 🔵 Improving AND Signal = 🟢 Accelerating → Priority B sectors (hidden gems)")
bullet("Note top 2–3 sector names. You will bias your stock scan toward these sectors.")
note("EXAMPLE",
     "If Nifty IT is Leading+Accelerating and Nifty Auto is Leading+Decelerating this week, "
     "scan IT stocks first and with more conviction. Auto stocks may be peaking.", color=C_AMBER)

h3("Step 3: Universe Scan with Beta Screener v2.3 (10 min)")
body("In TradingView Stock Screener, apply the Commander Screener Beta Edition v2.3 with these filters:")
code("Alpha Stars >= 4  AND  Stage Num = 2.1 (UP)  AND  Catalyst ID > 0")
code("AND  T1 Resistance Warning = 0  AND  Anti-Algo Gate = 1  AND  Vol Shelf = 1")
body("Sort by: Dashboard Quality Score (Descending)")
body("Result: Typically 5–15 stocks pass. These are your candidates for this week.")
body("Alternative: Commander Web → HUNTER → SCANNERS to run integrated scan within the app.")

h3("Step 4: Sector Filter & Prioritisation")
body("From the scanner results, prioritise stocks in your Priority A and B sectors from Step 2.")
body("If a stock is in a Priority A sector AND has Catalyst ID = 2 (POS-BO) AND SMC sweep confirmed → top priority.")

h2("26.3 PHASE 2: Fundamental Gate (15–20 minutes)")

h3("Step 1: Commander Web HUNTER → ENRICHMENT")
body("Load your top 10 candidates into the HUNTER ENRICHMENT tab:")
bullet("The app calls get_weinstein_score() for each symbol in parallel (ThreadPoolExecutor)")
bullet("Review AI Grade column: A or B = proceed. C = swing only. D/F = skip.")
bullet("Review Catalyst tag: POS-BO, SWG-BO, POS-ACCUM = highest priority")

h3("Step 2: Fundamental X-Ray Validation (for top 5 candidates)")
body("Load each candidate on TradingView with Fundamental X-Ray v2.2 active. Run the 15-second check:")
bullet("Overall Fundamental Rating: ≥ 9 (Grade B) for positional. ≥ 6 (Grade C) for swing.")
bullet("Piotroski Score: ≥ 5 for positional. ≥ 3 for swing.")
bullet("EPS Accelerating: YES → upsize by 25%. NO → standard size.")
bullet("Debt Trap: D/E > 1.5 AND India 10Y Yield RISING → remove from list.")
bullet("FCF: Negative FCF → swing only, no positional.")

h3("Step 3: Fundamental Veto Rules")
body("Remove from list if any of these are true:")
bullet("Overall Grade D or F (score < 6)")
bullet("Net Income < 0 (loss-making)")
bullet("D/E > 2.0 (highly over-leveraged)")
bullet("FCF negative AND Minervini Score < 4 (poor growth AND burning cash)")
bullet("Accrual Ratio > 10% (earnings quality concern)")

h2("26.4 PHASE 3: Structure Validation (10–15 min per candidate)")

h3("Step 1: Zigzag Structural Check")
body("Open the stock on TradingView Daily chart with Swing Zigzag v6.0 active:")
bullet("HUD: TREND = UPTREND — if SIDEWAYS or DOWNTREND, remove from list")
bullet("HUD: STRUCTURE = HH/HL — any other combination means trend is weakening")
bullet("HUD: CHOPPINESS ≤ 4 flips (GREEN) — if YELLOW or RED, skip this stock")
bullet("HUD: SWING COUNT = 1 or 2 (fresh trend) — ideal for breakout entries")
bullet("If SWING COUNT ≥ 4 — switch to pullback entry only, do not buy breakout")
bullet("Fib lines solid (not dotted) — confirms trend is actionable")

h3("Step 2: Dashboard v63.3 Full Analysis")
body("Load Dashboard v63.3 on the same chart and review:")
bullet("Stage Display: 🚀 STAGE 2 (Full Confluence) is ideal. ⚓ STAGE 2 (Value Zone) = pullback entry.")
bullet("Action Signal: ≥ 4/7 stars for a high-conviction setup")
bullet("Alpha Score: ≥ 60 (Grade B or better)")
bullet("Catalyst tag: Matches what Beta Screener showed — consistency across tools is confirmation")
bullet("Room for Trade: T1 not blocked by overhead resistance")
bullet("Sector Velocity: 🔥 ACCELERATING LEADER or ⚡ HIDDEN ACCUMULATION preferred")
bullet("Volatility State: SQUEEZE or VCP (True Dry Up) = ideal — energy coiled before breakout")
bullet("RS vs N500: LEADING (positive, rising slope) — must not be LAGGING")
bullet("Next Earnings: Not within 3 days of planned entry")

h2("26.5 PHASE 4: Execution (5–10 min per trade)")

h3("Step 1: Commander Web AI LAB → PRE-FLIGHT Check")
body("Run the 10-gate Pre-Flight check in Commander Web for the final candidate:")
body("All 10 gates must show GREEN. Any RED = do not enter today. Wait or skip.")

h3("Step 2: Strategy v4.5 Risk Allocator")
body("On TradingView, load Strategy v4.5 on the stock's daily chart:")
bullet("The Risk Allocator table shows: Entry, Stop, Quantity, T1, T2, Trade Validation")
bullet("Verify Trade Validation = '✅ Clear Room to T1' — if it shows T1 > Overhead Resistance, skip")
bullet("Note the QUANTITY — use this exact number in your broker order")
bullet("Note the Adjusted Risk % — Kelly has already factored setup quality")
bullet("If Distance to SL > 8% (from Beta Screener output) → halve the quantity manually")

h3("Step 3: Order Placement")
bullet("Place limit buy at the Entry price shown in Risk Allocator")
bullet("Set hard GTC stop loss at the Stop price shown in Risk Allocator")
bullet("Confirm order filled before logging the trade")

h3("Step 4: Log the Trade in Commander Web Journal")
body("Commander Web → JOURNAL → Add Entry:")
tbl(
    ["Field","What to Enter"],
    [
        ["Symbol","Exact NSE ticker (e.g., RELIANCE)"],
        ["Entry Date","Today's date"],
        ["Buy Price","Your actual fill price"],
        ["Stop Loss","SL from Strategy v4.5 Risk Allocator"],
        ["Target","T1 from Risk Allocator"],
        ["Quantity","Number of shares bought"],
        ["Sector","From X-Ray sector classification"],
        ["Catalyst","From Beta Screener (POS-BO, SWG-PB, etc.)"],
        ["Planned RR","2.5R (bull) or 2.0R (correction)"],
        ["Rationale","Stage, catalyst, key technical and fundamental factors"],
    ],
    col_widths=[2.0,5.0]
)

h2("26.6 PHASE 5: Active Trade Management (5–10 min daily)")

h3("Daily Pre-Market Check (3:10 PM IST or next-morning pre-open)")
bullet("Open Commander Web → DASHBOARD: Check Market Status — any change from yesterday?")
bullet("If market degraded (Bull → Correction → Bear): Tighten ALL stops immediately")
bullet("Open Commander Web → COMMAND → ACTIVE OPS: Check for any ⚠ STOP HIT or Noise flags")
bullet("Open Dashboard v63.3 for any positions showing STOP HIT: Place manual exit order for open")

h3("Intraday Monitoring (Optional, not required)")
bullet("Dashboard TRADE STATUS column: TRIM or SCALE IN requires action")
bullet("TRIM: Sell the indicated % at market — do not delay profit-taking")
bullet("SCALE IN: Only if initial position is at breakeven or above (position is 'free')")

h3("Weekly Position Review (Every Sunday)")
bullet("Step 1: Recalculate Chandelier Exit for all positional trades: Highest High (14d) − 3.0×ATR14")
bullet("Step 2: Update p_sl in Dashboard v63.3 (drag the orange stop line) if new CE > old SL")
bullet("Step 3: Check Days Held vs Time Stop: any position approaching 6-week limit with < 0.5R profit?")
bullet("Step 4: Check Zigzag Swing Count on all positions: ≥ 4 swings = consider 50% profit taking")
bullet("Step 5: Review Sector Velocity for all open positions: any EXHAUSTED LEADER showing?")
bullet("Step 6: Run Commander Web AUTOPSY if ≥ 5 trades closed this month")

h2("26.7 Entry Decision Matrix (Full Flowchart)")
tbl(
    ["Question","YES","NO"],
    [
        ["Is Market State = BULLISH?","Continue","Reduce size to 75%. Only STRONG BUY signals."],
        ["Is Stage = Stage 2 (any variant)?","Continue","Stop. No entry. Watch if Stage 1 Constructive."],
        ["Is Choppiness ≤ 4 flips (Zigzag)?","Continue","Stop. Find a different stock."],
        ["Is RS vs CNX500 Leading or Improving?","Continue","Skip unless exceptional other factors."],
        ["Is Overhead Resistance ≤ 3 levels?","Continue","Likely BLOCKED — confirm Trade Validation."],
        ["Is Catalyst detected (non-NONE)?","Continue","Add to watchlist. Set BoS alert."],
        ["Is Trade Validation = Clear Room to T1?","Continue","Skip. Bad RR. Wait for better entry."],
        ["Is Fundamental Grade ≥ B (positional) or ≥ C (swing)?","Continue","Swing only at C. Skip at D/F."],
        ["Is EPS Accelerating?","Upsize 25%","Standard size."],
        ["Are all 10 Pre-Flight gates GREEN?","EXECUTE","Wait. One red gate = no trade."],
    ],
    col_widths=[3.5,2.0,2.5]
)

divider()

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 27 — RISK MANAGEMENT FRAMEWORK
# ═══════════════════════════════════════════════════════════════════════════
h1("Section 27 — Bull Market Risk Management Framework")

h2("27.1 Core Position Sizing Formula")
code("Risk Amount   = Portfolio Capital × Adjusted Risk %")
code("Risk Per Share = Entry Price − Stop Loss Price")
code("Risk-Sized Qty = Risk Amount / Risk Per Share")
code("Cap-Sized Qty  = Max Allocation / Entry Price")
code("Final Quantity = MIN(Risk-Sized Qty, Cap-Sized Qty)")
body("Example:")
code("Portfolio = ₹10,00,000 | Stock Risk = 0.75% → Risk Amount = ₹7,500")
code("Entry = ₹500, SL = ₹475 → Risk Per Share = ₹25")
code("Risk-Sized Qty = 7,500 / 25 = 300 shares")
code("Cap-Sized Qty  = 25,000 / 500 = 50 shares")
code("Final Qty = MIN(300, 50) = 50 shares | Invested = ₹25,000")

h2("27.2 Risk Percentage by Asset Type and Regime")
tbl(
    ["Asset Type","Bull Market","Correction","Bear Market"],
    [
        ["ETF (index/sector)","1.0%","0.75%","0.5%"],
        ["Stock (equity)","0.75%","0.5%","0.25%"],
        ["Stock (EPS Accelerating)","0.94% (1.25× Kelly)","0.63%","0.31%"],
        ["Stock (weak setup)","0.56% (0.75× Kelly)","0.38%","0.19%"],
        ["Absolute Floor","0.25%","0.25%","0.25%"],
    ],
    col_widths=[2.5,1.5,1.5,1.5]
)

h2("27.3 Portfolio Exposure Limits")
tbl(
    ["Rule","Bull Market","Correction","Bear Market"],
    [
        ["Maximum positions open","6","4","0 (exits only)"],
        ["Maximum single position","25% of portfolio","20% of portfolio","N/A"],
        ["Maximum sector exposure","30% of portfolio","20% of portfolio","N/A"],
        ["Maximum total long exposure","80% of portfolio","50% of portfolio","Cash ≥ 50%"],
        ["Correlated positions (r > 0.85)","Max 1 shadow pair","None","N/A"],
    ],
    col_widths=[3.0,1.8,1.8,1.8]
)

h2("27.4 Stop Loss Levels by Catalyst Type")
tbl(
    ["Catalyst","Initial Stop Placement","Trail Method","Emergency Exit"],
    [
        ["POS-BO","10-bar structural low − 0.2×ATR14","Chandelier 3.0×ATR (tightens at 1.5R and 3.0R)","Stage 4 confirmed"],
        ["POS-ACCUM","10-bar structural low − 0.2×ATR14","Chandelier 3.0×ATR","Stage 4 confirmed"],
        ["SWG-PB","EMA20 × (1−1%)","EMA20 × (1−1%) trailing","Close below SMA50"],
        ["SWG-BO","5-bar low − 0.2% padding","EMA20 × (1−1%) trailing","Close below SMA50"],
        ["SWG-GAP","Gap bar low − 0.2% padding","EMA20 × (1−1%) trailing","Close below SMA50"],
        ["SWG-REV","5-bar low − 0.2% padding","5-day hard exit (no trail)","5-day time stop"],
    ],
    col_widths=[1.5,2.2,2.5,2.0]
)

h2("27.5 Adaptive Chandelier Exit Levels")
tbl(
    ["Profit Level","Multiplier (Bull)","Multiplier (Bear)","Effect"],
    [
        ["0 → 1.5R","3.0×ATR14","3.5×ATR14","Maximum room for trade to breathe"],
        ["1.5R → 3.0R","2.25×ATR14 (3.0×0.75)","2.625×ATR14","Tightening — protecting accumulated profit"],
        ["3.0R+","1.5×ATR14 (3.0×0.50)","1.75×ATR14","Super-tight — maximum profit capture"],
    ],
    col_widths=[1.8,1.8,1.8,3.1]
)

h2("27.6 Target Levels — Dynamic R:R")
tbl(
    ["Trade Type","T1 (Partial Exit)","% to Sell at T1","T2 (Runner Exit)","% to Sell at T2","Final Runner"],
    [
        ["POS-BO / POS-ACCUM (Bull)","2.5R","30% of position","Managed by CE trail","—","70% of original — trail with Chandelier"],
        ["SWG-PB / SWG-BO / SWG-GAP (Bull)","2.5R","50% of position","3.5R","50% of remaining (25% original)","25% of original — trail with EMA20"],
        ["SWG-REV (always)","2.0R","100% — full exit at T1","—","—","No runner for mean reversion"],
        ["Any setup (Correction)","2.0R","50% of position","3.0R","50% of remaining","25% trail with EMA20"],
    ],
    col_widths=[2.5,0.8,1.5,0.8,1.5,2.4]
)
note("RULE",
     "At T1: Always move stop to breakeven on the remaining position. "
     "A trade that reaches T1 should NEVER become a losing trade. "
     "This is the most important rule in trade management.", color=C_RED)

h2("27.7 Noise Risk Protocol")
body(
    "A 'noisy' stop loss is one placed closer to the entry than 1.5× ATR14 — "
    "normal intraday volatility will trigger it before the trade develops."
)
body("Commander Web DASHBOARD shows a noise risk badge: ⚠ N AT RISK")
body("If your position is flagged:")
bullet("Option 1: Widen stop to structural low (proper placement) and reduce size proportionally")
bullet("Option 2: Accept the risk but acknowledge it is a 'micro-stop' trade")
bullet("Option 3: Exit the position and re-enter at a proper technical stop level")
note("NEVER", "Do NOT simply widen the stop without reducing position size. "
     "Widening a stop without reducing size increases your rupee risk.", color=C_RED)

h2("27.8 Portfolio Correlation Rules")
body(
    "Commander Web COMMAND page uses get_portfolio_correlation_matrix() to identify "
    "'shadow pairs' — positions with correlation r > 0.85."
)
bullet("Shadow pairs move together — when one stops out, the other likely stops out within 1–2 days")
bullet("Maximum: 1 shadow pair at any time in a bull market. Zero in correction.")
bullet("Target portfolio diversity score > 0.6 (output from correlation matrix)")
bullet("If adding a new position would create a second shadow pair: choose a different stock or sector")

divider()

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 28 — COMMON MISTAKES & HOW TO AVOID THEM
# ═══════════════════════════════════════════════════════════════════════════
h1("Section 28 — Common Mistakes & How to Avoid Them")

tbl(
    ["Mistake","What Goes Wrong","Prevention Rule"],
    [
        ["Buying Stage 4 hoping for reversal",
         "Stage 4 stocks continue lower 70%+ of the time. 'Value buying' in decline phases destroys capital.",
         "Stage 2 only. The state machine must confirm Stage 2 before any entry is permitted."],
        ["Ignoring the Trade Validation warning",
         "T1 is above the 52-week high or major resistance. Price reverses before reaching T1 and you exit at a loss.",
         "Only enter when Trade Validation = 'Clear Room to T1'. If BLOCKED, wait for price to break above resistance and re-base."],
        ["Trading with RS Lagging",
         "Even in bull markets, laggard stocks underperform. You make 5% on a stock while the market makes 15%.",
         "RS vs CNX500 must be at minimum 'Improving'. Never enter when RS = 'Lagging'."],
        ["Chasing gaps without Anti-Algo check",
         "HFT algorithms spike price above resistance intrabar, triggering buy signals, then reverse. You buy the spike top.",
         "Anti-Algo Gate MUST = 1 for all SWG-BO and SWG-GAP entries. The close must be in top 40% of bar."],
        ["Over-sizing in correction markets",
         "Normal 3–5% corrections hit stops faster. Losing 3 in a row at 0.75% each = 2.25% drawdown, not alarming. At 1.5% each = 4.5% sudden drawdown.",
         "Correction mode: reduce risk % to 0.5% for stocks. Kelly multiplier automatically derates weak setups."],
        ["Holding Stage 3 positions",
         "Distribution phases are violent. Stocks fall 20–40% rapidly from Stage 3 peaks. The greed of 'waiting for one more high' is catastrophic.",
         "Zigzag Bearish BoS alert triggers on breakdown — exit immediately on Stage 3 detection."],
        ["Not enforcing time stops",
         "Capital is tied up in stagnant trades for weeks. The opportunity cost is massive — other stocks make 30% while yours goes sideways.",
         "Strategy v4.5 auto-exits at 6 weeks (positional) / 10 days (swing) if < 0.5R profit. Mirror this in your broker."],
        ["Using wrong sector for RS",
         "Bank stock classified as 'Finance' instead of 'Banking' gets incorrect sector RS. Recommendation shows incorrectly.",
         "Enable Auto-Detect Sector in Strategy v4.5. If sector detection fails, manually override with correct sector index."],
        ["Fundamentally weak fast-movers",
         "Pump-and-dump stocks make spectacular moves then collapse 50–80%. No institutional support = no sustainable trend.",
         "X-Ray Fundamental Grade must be ≥ C for swing. ≥ B for positional. F-rated stocks = scalp only if at all."],
        ["Moving stop down to 'give the trade more room'",
         "Moving a stop down after a loss violates the original risk thesis. The initial stop was placed at the structural invalidation level.",
         "Stops ONLY move up. Never backward. If the original stop level is hit, the thesis is wrong — exit."],
        ["Entering SWG-REV in Stage 4",
         "Before v4.5 fix: RSI-3 < 20 could fire in Stage 4 (dead-cat bounce in downtrend). Result: immediate continuation lower.",
         "v4.5 fixes this. SWG-REV now requires Stage 1 OR Stage 2. If using v4.4 or earlier, manually verify Stage first."],
        ["Ignoring the Choppiness reading",
         "High-choppiness stocks (>8 flips/52W) have erratic price behaviour. Normal stop placements are systematically hunted.",
         "Choppiness > 8 (red) on Zigzag = remove stock from consideration immediately. No exceptions."],
        ["Skipping the 10-gate Pre-Flight check",
         "One gate failure (e.g., earnings in 2 days, or shadow pair limit hit) means the setup has known risk that you are choosing to ignore.",
         "Run Commander Web AI LAB → PRE-FLIGHT for every new trade. All 10 gates GREEN = proceed. 9/10 = no trade."],
        ["Over-riding the Kelly size to bet bigger",
         "Kelly says 0.75× (weak setup). You override to 1.25× because you 'feel' the trade is good. You lose 1.25× on a 0.75× quality setup.",
         "Trust the Kelly output. The system has assessed Stage, RS, and Volume objectively. Your feelings have not."],
    ],
    col_widths=[2.0,2.8,3.2]
)

divider()

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 29 — PSYCHOLOGY: THE TRADER'S MINDSET
# ═══════════════════════════════════════════════════════════════════════════
h1("Section 29 — The Trader's Mindset: System Maxims")

h2("29.1 The Five Maxims of This System")

h3("1. Do Not Override the Stop Loss Calculation")
body(
    "The Strategy v4.5 dynamically calculates stops based on ATR volatility and structural lows "
    "to give the trade exact breathing room without exposing tail risk. "
    "When you move a stop down to 'give the trade more room,' you are replacing a mathematical "
    "decision with an emotional one. The math wins long-term. Your feelings do not."
)

h3("2. Respect the Fundamental Filter")
body(
    "In bull markets, even garbage stocks fly. Fundamentally broken companies make spectacular "
    "50–100% moves. It is tempting to ride them. But when the market turns — and it always does — "
    "these stocks gap down 20–40% with no buyers. The X-Ray module and Alpha Scoring engines "
    "protect your account from ruin. Use them."
)

h3("3. Volume Is Truth")
body(
    "Breakouts without volume expansion are suspect. Pullbacks without volume dry-up are suspect. "
    "All active modules in this ecosystem require relVol > 1.25× for breakouts and < 0.70× "
    "for pullbacks. These are coded thresholds backed by backtesting. "
    "A beautiful chart pattern with low volume at the pivot is not a signal — it is a trap."
)

h3("4. Trade the Math, Not the News")
body(
    "If a stock releases exceptional earnings but simultaneously drops into Stage 4 with a "
    "Bearish BoS on the Zigzag — you sell. The story does not matter. "
    "Institutional capital is always revealed in the 30-Week MA direction and OBV trajectory "
    "before the fundamental story reaches retail investors. "
    "When the chart says sell, you sell. When it says wait, you wait."
)

h3("5. Capital Velocity Over Capital Preservation of Losers")
body(
    "A stagnant trade is not a safe trade — it is an opportunity cost. "
    "Time stops (6 weeks positional, 10 days swing) enforce capital velocity. "
    "Dead money in a stagnant trade is capital that is not compounding in a moving trade. "
    "Accept small losses and time-stop exits without guilt — they are not failures, "
    "they are the system working correctly."
)

h2("29.2 The Trader's Weekly Commitment")
body("Every Sunday, commit to these three things:")
bullet("1. Follow the system — run all five phases. Do not skip the fundamental check because 'I know this company.'")
bullet("2. Update all trailing stops — do not let a winner turn into a loser by neglecting the weekly SL update.")
bullet("3. Review performance honestly — if expectancy is negative over the last 10 trades, find one specific gate that was violated most. Fix that gate.")

h2("29.3 Position Sizing Psychology")
body(
    "The hardest trade to size correctly is the one you are most excited about. "
    "Excitement triggers overconfidence — you increase size, the trade fails, you lose more than planned. "
    "The Kelly Fraction was designed for this: it sizes you smaller when setup quality is lower "
    "and larger only when the objective algorithm confirms a high-probability setup. "
    "Let the algorithm be your emotion filter."
)

h2("29.4 On Losses")
body(
    "A trade that hits its stop loss and exits at the planned SL is NOT a bad trade — "
    "it is a perfect trade execution. The stop was placed correctly, the system worked, "
    "you lost exactly what you planned to lose. "
    "A bad trade is one that runs past its stop because you hesitated. "
    "The goal is not to avoid losses. The goal is to make your losses SMALL and PLANNED."
)

divider()

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 30 — BULL MARKET QUICK REFERENCE CARD
# ═══════════════════════════════════════════════════════════════════════════
h1("Section 30 — Bull Market Quick Reference Card")

h2("Pre-Trade Checklist (Print & Use Daily)")
tbl(
    ["#","Gate","Condition","Tool"],
    [
        ["1","Market State","BULLISH or CORRECTION (never BEARISH)","Commander Web DASHBOARD"],
        ["2","Stage","Stage 2 (any variant) on weekly chart","Dashboard v63.3 / Ultimate v3.6"],
        ["3","Zigzag Structure","UPTREND + HH/HL + Choppiness ≤ 4","Swing Zigzag v6.0"],
        ["4","RS vs CNX500","Leading or Improving (not Lagging)","Dashboard / Ultimate"],
        ["5","Overhead Resistance","≤ 3 levels OR Trade Validation = Clear","Strategy v4.5 Risk Allocator"],
        ["6","Catalyst","Non-NONE (active catalyst detected)","Beta Screener / Dashboard"],
        ["7","Anti-Algo Gate","= 1 for SWG-BO and SWG-GAP","Beta Screener v2.3"],
        ["8","Fundamental Grade","≥ B for positional / ≥ C for swing","X-Ray v2.2"],
        ["9","Earnings Proximity","Not within 3 days","Dashboard v63.3"],
        ["10","All 10 Pre-Flight Gates","All GREEN","Commander Web AI LAB"],
    ],
    col_widths=[0.4,1.8,2.8,2.5]
)

h2("Catalyst Quick-Reference")
tbl(
    ["Catalyst","Type","Hold Period","T1","T2","Size"],
    [
        ["POS-BO","Positional Breakout","Weeks–months","2.5R (30% exit)","CE trail","Standard (Kelly-adjusted)"],
        ["POS-ACCUM","Positional (early)","Weeks–months","2.5R (30% exit)","CE trail","Standard"],
        ["SWG-PB","Swing Pullback","1–3 weeks","2.5R (50% exit)","3.5R (25% exit)","Standard"],
        ["SWG-BO","Swing VCP Break","1–2 weeks","2.5R (50% exit)","3.5R (25% exit)","Standard"],
        ["SWG-GAP","Gap & Go","1–10 days","2.5R (50% exit)","3.5R (25% exit)","Standard"],
        ["SWG-REV","Mean Reversion","1–5 days (hard stop)","2.0R (100% exit)","None","50–75% of standard"],
    ],
    col_widths=[1.5,1.8,1.5,1.5,1.5,1.7]
)

h2("Trade Status Actions (Dashboard v63.3)")
tbl(
    ["Dashboard Shows","Required Action","Urgency"],
    [
        ["⚠️ STOP HIT (P&L%)", "Exit immediately at next market open. No hesitation.", "IMMEDIATE"],
        ["EXIT NOW (40WMA Break)", "Close full position regardless of trailing stop level.", "IMMEDIATE"],
        ["TRIM (P&L%)", "Sell the indicated % of position at current market price.", "Same session"],
        ["SCALE IN (P&L%)", "Add up to 50% of original size only if initial stop ≥ breakeven.", "Optional"],
        ["HOLD (P&L%)", "No action. Review weekly trailing stop level.", "Weekly"],
        ["🔻 STAGE 4 (Liquidate)", "Exit entire position immediately. Institutional trend ended.", "IMMEDIATE"],
        ["⏰ POS TIME STOP HIT", "Close position at market. 6 weeks with < 0.5R = dead money.", "Same session"],
        ["⏰ SWG TIME STOP HIT", "Close position at market. 10 days with < 0.5R = dead money.", "Same session"],
    ],
    col_widths=[2.5,3.5,1.5]
)

h2("Module Stack for Each Candidate Stock")
tbl(
    ["Module","Pane","Purpose"],
    [
        ["Swing Pro Dashboard v63.3","Overlay on price","Primary analysis — all metrics, portfolio tracking"],
        ["Swing Zigzag Strict v6.0","Overlay on price","Structure visualisation — HH/HL, Fib levels, BoS alerts"],
        ["Fundamental X-Ray v2.2","Overlay (bottom-left)","Quality validation — Minervini + Piotroski + 17-pt rating"],
        ["Strategy v4.5","Overlay on price","Entry signals + Risk Allocator table + trailing stop lines"],
    ],
    col_widths=[2.5,1.5,4.0]
)

h2("Key Formulas (All Modules)")
for line in [
    "Mansfield RS        = (RS_Line / SMA26_of_RS_Line − 1) × 10",
    "Risk Amount         = Portfolio × Risk% (0.75% stocks / 1.0% ETF)",
    "Position Size       = Risk Amount / (Entry − Stop Loss)",
    "Chandelier Exit     = Highest High (14 bars) − 3.0 × ATR14 (Bull)",
    "Kelly High-Prob     = base_risk × 1.25 (Stage2 + RS + Vol all pass)",
    "T1 Positional       = Entry + (Entry − SL) × 2.5",
    "T2 Swing            = Entry + (Entry − SL) × 3.5 (Bull)",
    "Algo SL Positional  = 10-bar structural low − ATR14 × 0.2",
    "VCP Tight           = ATR10 < SMA50(ATR10) × 1.5",
    "Vol Dry-Up          = Volume < SMA50(Volume) × 0.70",
    "Anti-Algo Gate      = (close − low) / (high − low) ≥ 0.60",
    "SMC Sweep           = 5-bar low < SMA50 AND close > SMA50 AND vol > SMA50(vol) × 1.5",
    "Weekly Slope Proxy  = (SMA150 − SMA150[20]) / 20.0",
]:
    code(line)

h2("NSE Sector Index Reference")
tbl(
    ["Sector","TradingView Symbol","Use For"],
    [
        ["Banking & Finance","NSE:BANKNIFTY","Banks, NBFCs, Insurance companies"],
        ["Information Technology","NSE:CNXIT","IT services, software, tech"],
        ["Pharmaceuticals","NSE:CNXPHARMA","Pharma, healthcare, diagnostics"],
        ["Automobiles","NSE:CNXAUTO","Auto OEMs, components, EV"],
        ["FMCG & Consumer","NSE:CNXFMCG","Consumer staples, food, FMCG"],
        ["Metals & Mining","NSE:CNXMETAL","Steel, aluminium, mining"],
        ["Energy","NSE:CNXENERGY","Oil & gas, power, renewables"],
        ["Real Estate","NSE:CNXREALTY","Builders, REITs, property"],
        ["Infrastructure","NSE:CNXINFRA","Capital goods, utilities, construction"],
        ["Broad Market (benchmark)","^CRSLDX (CNX500)","Default RS benchmark for all calculations"],
    ],
    col_widths=[2.0,2.0,3.5]
)

doc.add_page_break()
doc.save(OUTPUT)
print(f"Part 7 saved: {OUTPUT}")
print("Sections 26–30 complete. Bull Market section DONE.")
