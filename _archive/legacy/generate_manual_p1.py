"""
Recovery Trading Manual — Generator Part 1
Sections 1–8: Strategy overview through entry rules
Run this first to create the base document.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUTPUT = r"C:\Users\jayra\Documents\GeminiVSCode\Recovery_Trading_Manual.docx"

# ── Colour palette ──────────────────────────────────────────────────────────
C_NAVY   = RGBColor(0x1F, 0x35, 0x64)   # section headers
C_TEAL   = RGBColor(0x00, 0x70, 0x70)   # sub-section headers
C_AMBER  = RGBColor(0xC5, 0x5A, 0x11)   # warnings / stop-loss
C_GREEN  = RGBColor(0x37, 0x86, 0x30)   # buy / positive signals
C_LGREY  = RGBColor(0xF2, 0xF2, 0xF2)   # table shading
C_DGREY  = RGBColor(0x40, 0x40, 0x40)   # body text default
C_RED    = RGBColor(0xC0, 0x00, 0x00)   # sell / danger

doc = Document()

# ── Page setup ───────────────────────────────────────────────────────────────
for section in doc.sections:
    section.page_width  = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = section.right_margin  = Inches(1.0)
    section.top_margin  = section.bottom_margin = Inches(0.9)

# ── Style helpers ────────────────────────────────────────────────────────────
def set_font(run, bold=False, italic=False, size=11, color=C_DGREY):
    run.bold   = bold
    run.italic = italic
    run.font.size  = Pt(size)
    run.font.color.rgb = color

def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    set_font(run, bold=True, size=18, color=C_NAVY)
    return p

def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    set_font(run, bold=True, size=14, color=C_TEAL)
    return p

def h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    set_font(run, bold=True, size=12, color=C_NAVY)
    return p

def body(text, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(indent * 0.25)
    run = p.add_run(text)
    set_font(run, size=10.5)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent  = Inches(0.25 + level * 0.2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    set_font(run, size=10.5)
    return p

def note(label, text, color=C_AMBER):
    p = doc.add_paragraph()
    r1 = p.add_run(f"{label}: ")
    set_font(r1, bold=True, size=10.5, color=color)
    r2 = p.add_run(text)
    set_font(r2, size=10.5)
    return p

def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run("─" * 85)
    set_font(run, size=8, color=RGBColor(0xCC, 0xCC, 0xCC))

def shade_cell(cell, rgb: RGBColor):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:fill'), f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}")
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)

def simple_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        shade_cell(cell, C_NAVY)
        p = cell.paragraphs[0]
        run = p.add_run(h)
        set_font(run, bold=True, size=10, color=RGBColor(0xFF,0xFF,0xFF))
    for ri, row_data in enumerate(rows):
        row = t.add_row()
        bg = C_LGREY if ri % 2 == 0 else RGBColor(0xFF, 0xFF, 0xFF)
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            shade_cell(cell, bg)
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            set_font(run, size=10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t

# ═══════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
run = p.add_run("RECOVERY TRADING MANUAL")
set_font(run, bold=True, size=28, color=C_NAVY)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Weinstein Stage Analysis · Phase-1 Recovery Trades")
set_font(run, italic=True, size=14, color=C_TEAL)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Weinstein Commander Suite v3.0  ·  Dashboard v63.3  ·  Recovery Strategy v1.0")
set_font(run, size=11, color=C_DGREY)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Internal Reference — Not for Distribution")
set_font(run, italic=True, size=10, color=RGBColor(0x80,0x80,0x80))

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 1 — PHILOSOPHY
# ═══════════════════════════════════════════════════════════════════════════
h1("Section 1 — Philosophy & Core Principles")

h2("1.1 What Is a Recovery Trade?")
body("A recovery trade targets high-quality companies temporarily dislocated from their primary uptrend. "
     "The goal is to buy before the institutional crowd recognises the repair. "
     "This is NOT distressed investing — we only trade companies with positive earnings, free cash flow, "
     "and healthy balance sheets that are stuck in a broader market correction or sector rotation drawdown.")

h2("1.2 The Three Recovery Edges")
simple_table(
    ["Edge", "Code", "Setup Description", "Risk Level"],
    [
        ["Climax Bottom Bounce", "REV-CB", "Stock stretched ≥8% below 30W SMA with ≥2× volume climax. RS line positive. Bounce entry.", "High — only T1 target"],
        ["RS Survivor Breakout", "REV-RS", "RS line already positive vs CNX500 while price still below 200d SMA. Volume-confirmed breakout.", "Medium — can run to T2"],
        ["Early VCP / Golden Cross", "REV-EARLY", "SMA50 ≥92% of SMA200 (near golden cross). VCP tightening. RS positive. Pre-breakout entry.", "Low — highest quality"],
    ],
    col_widths=[1.5, 0.8, 3.5, 1.2]
)

h2("1.3 Why Market Regime Matters")
body("Recovery trades only fire when the market regime allows. The Three-Way Regime Gate blocks entries when:")
bullet("CNX500 is in a death cross (50d < 200d SMA) AND")
bullet("CNX500 is ≥7% off its 52-week high, OR")
bullet("The individual stock is ≥10% off its own 52-week high")
note("RULE", "Never force a recovery trade into a bear market. The regime gate is non-negotiable.")

h2("1.4 Position Sizing Philosophy")
body("Recovery trades use asymmetric sizing. You risk less per trade (0.5–1.0% of capital) but allow more room "
     "for the setup to develop. The wider Chandelier Exit (3.5× ATR14) reflects recovery-phase volatility.")

divider()

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 2 — WEINSTEIN STAGE ANALYSIS
# ═══════════════════════════════════════════════════════════════════════════
h1("Section 2 — Weinstein Stage Analysis Primer")

h2("2.1 The Four Stages")
simple_table(
    ["Stage", "30W SMA", "Price Position", "Volume", "Action"],
    [
        ["Stage 1 — Base", "Flat", "Near/at 30W SMA", "Contracting", "Watch only"],
        ["Stage 2 — Advance", "Rising", "Above 30W SMA", "Expanding on up-days", "BUY / HOLD"],
        ["Stage 3 — Top", "Flat/rolling", "Choppy around 30W SMA", "Heavy, two-sided", "Reduce / exit"],
        ["Stage 4 — Decline", "Falling", "Below 30W SMA", "Expanding on down-days", "AVOID / short"],
    ],
    col_widths=[1.5, 1.2, 2.0, 2.0, 1.3]
)

h2("2.2 Recovery Phase Definition")
body("Recovery trades occur at the transition from Stage 4 → Stage 1 → early Stage 2. "
     "Key signals: 30W SMA slope begins to flatten after a decline, price forms a higher low, "
     "RS line starts rising relative to CNX500.")

h2("2.3 Stage Identification in Dashboard v63.3")
body("The Dashboard displays stage with emoji labels:")
bullet("🚀 STAGE 2 (Full Confluence) — All Minervini criteria met")
bullet("⚓ STAGE 2 (Value Zone) — Stage 2 but price pulled back to 20 EMA")
bullet("👀 STAGE 1 (Constructive) — Base forming, 30W SMA flattening")
bullet("😴 STAGE 1 (Dormant) — Base with no RS/volume confirmation")
bullet("🛑 STAGE 3 (Volatile Churn) — Top formation, avoid")
bullet("🔻 STAGE 4 (Liquidate) — Confirmed decline, exit all longs")

h2("2.4 Mansfield RS Formula")
body("Relative Strength is measured on a weekly basis against CNX500 (^CRSLDX):")
p = doc.add_paragraph()
run = p.add_run("    RS = (RS_Line / SMA26_of_RS_Line − 1) × 10")
set_font(run, bold=True, size=11, color=C_NAVY)
body("A positive RS value means the stock is outperforming CNX500 on a momentum-adjusted basis. "
     "Recovery setups require RS > 0. RS ≥ 1.0 earns an extra point in the composite score.")

divider()

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 3 — RECOVERY FUNDAMENTAL FILTER (RFF)
# ═══════════════════════════════════════════════════════════════════════════
h1("Section 3 — Recovery Fundamental Filter (RFF)")

h2("3.1 The Six Binary Checks")
body("The RFF is a pass/fail filter applied before any recovery trade is taken. "
     "Each check is binary (pass=1, fail=0). Minimum score for a valid trade: 1/6 "
     "(configurable via rff_min_score in recovery_screener.py).")

simple_table(
    ["Check", "Metric", "Threshold", "Rationale"],
    [
        ["NI", "Net Income (TTM)", "> 0", "Company must be profitable"],
        ["FCF", "Free Cash Flow", "> 0", "Cash generation, not just accounting profit"],
        ["ICR", "Interest Coverage Ratio", "≥ 2×", "Debt service is sustainable"],
        ["DE", "Debt/Equity Ratio", "< 2", "Balance sheet not over-leveraged"],
        ["CR", "Current Ratio", "> 1", "Short-term liquidity adequate"],
        ["ROA", "Return on Assets", "> 5%", "Capital efficiency minimum"],
    ],
    col_widths=[0.6, 2.0, 1.2, 3.2]
)

h2("3.2 RFF in the Screener")
body("rff_min_score=1 means at least one of the six checks must pass. "
     "In practice, raise this to 3–4 for higher-conviction trades. "
     "The RFF score (0–6) contributes directly to the composite score.")

note("WARNING", "A stock that fails all 6 RFF checks is excluded regardless of technical setup.", color=C_RED)

divider()

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 4 — REGIME GATE
# ═══════════════════════════════════════════════════════════════════════════
h1("Section 4 — The Three-Way Regime Gate")

h2("4.1 Gate Logic")
body("All three conditions must be TRUE simultaneously to BLOCK a recovery signal:")
p = doc.add_paragraph()
run = p.add_run("    BLOCKED = (CNX500 death cross) AND (CNX500 ≥7% below 52W high) AND (Stock ≥10% below 52W high)")
set_font(run, bold=True, size=10.5, color=C_AMBER)

body("If any one condition is FALSE, the gate is OPEN and signals can fire.")

h2("4.2 Regime Contribution to Composite Score")
simple_table(
    ["Regime Condition", "Score Contribution"],
    [
        ["CNX500 above 200d SMA (healthy)", "+2 points"],
        ["CNX500 below 200d SMA but no death cross", "+1 point"],
        ["Full bear (death cross + correction)", "+0 points"],
    ],
    col_widths=[3.5, 2.5]
)

h2("4.3 Dashboard Market Health Display")
body("The Dashboard shows CNX500 health as one of four states:")
bullet("BULLISH (Confirmed) — CNX500 above all MAs, positive trend")
bullet("CORRECTION (In Uptrend) — Pullback within bull market")
bullet("RECOVERY (Weak) — Early signs of base building after bear")
bullet("BEARISH — Death cross confirmed, avoid new longs")

divider()

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 5 — THE COMPOSITE SCORE
# ═══════════════════════════════════════════════════════════════════════════
h1("Section 5 — Recovery Composite Score (0–20)")

h2("5.1 Score Components")
simple_table(
    ["Component", "Max Points", "Condition"],
    [
        ["RFF Score", "6", "0–6 fundamental checks passed"],
        ["RS Positive", "1", "Mansfield RS > 0"],
        ["RS Strong", "1", "Mansfield RS ≥ 1.0"],
        ["Discount Depth", "1–3", "1pt: 10–20% off high | 2pt: 20–30% | 3pt: >30%"],
        ["Regime", "1–2", "1pt: CNX500 below 200d | 2pt: CNX500 above 200d"],
        ["Stage", "1", "Price above 200d SMA (Stage 1/2 recovery)"],
        ["Edge Fired", "1–3", "1pt: CB-Watch | 2pt: REV-CB | 3pt: REV-RS or REV-EARLY"],
        ["Volume", "1–2", "1pt: above avg | 2pt: institutional surge (≥1.5×)"],
        ["Intersection Bonus", "1", "Multiple edges fire simultaneously"],
    ],
    col_widths=[2.0, 1.2, 3.8]
)

h2("5.2 Score Interpretation")
simple_table(
    ["Score", "Quality", "Action"],
    [
        ["16–20", "Elite setup", "Full position, prioritise"],
        ["12–15", "High conviction", "Standard position"],
        ["8–11", "Moderate", "Half position, wait for confirmation"],
        ["4–7", "Low conviction", "Watchlist only"],
        ["0–3", "Reject", "Do not trade"],
    ],
    col_widths=[1.0, 2.0, 3.5]
)

divider()

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 6 — ENTRY RULES BY EDGE
# ═══════════════════════════════════════════════════════════════════════════
h1("Section 6 — Entry Rules by Edge Type")

h2("6.1 REV-CB — Climax Bottom Bounce")
body("Best used for quick mean-reversion trades back toward the 30W SMA. Target only T1.")

simple_table(
    ["Rule", "Requirement"],
    [
        ["Price location", "≥8% below 30W SMA (cb_stretch_pct=8.0)"],
        ["Volume climax", "Climax candle volume ≥2× 50-day average (cb_vol_mult=2.0)"],
        ["Lookback window", "Climax must occur within last 10 bars (cb_climax_window=10)"],
        ["RS condition", "Mansfield RS > 0 (stock outperforming CNX500)"],
        ["Regime gate", "Must pass Three-Way Regime Gate"],
        ["RFF", "Minimum rff_min_score (default 1) checks passed"],
        ["Entry", "Buy next-day open or on first green candle after climax"],
        ["Stop", "Below climax low − 0.5×ATR14"],
        ["Target", "T1 only — 30W SMA level (2.0R)"],
    ],
    col_widths=[2.0, 5.0]
)
note("RULE", "Do NOT hold REV-CB for T2. This is a bounce trade. Exit at T1 regardless.")

h2("6.2 REV-RS — RS Survivor Breakout")
body("RS line is already positive while price is still recovering. Highest-probability multi-week hold.")

simple_table(
    ["Rule", "Requirement"],
    [
        ["RS condition", "RS Line > 30W SMA of RS Line (Mansfield RS > 0)"],
        ["Price location", "Below 200d SMA (still in recovery, not yet Stage 2)"],
        ["Above SMA50", "Price above 50d SMA (early recovery, not deeply broken)"],
        ["Weekly RSI", "> 55 (momentum building)"],
        ["Price minimum", "> ₹100 (liquidity filter)"],
        ["Volume", "Breakout candle ≥1.5× average (vol_confirm_mult=1.5)"],
        ["Breakout", "Price closes above 20-bar high on above-average volume"],
        ["Stop", "Chandelier Exit: 3.5 × ATR14 below recent pivot low"],
        ["Targets", "T1 at 2.0R, T2 at 2.5R"],
    ],
    col_widths=[2.0, 5.0]
)

h2("6.3 REV-EARLY — VCP Near Golden Cross")
body("Highest quality recovery setup. Stock is tightening in a VCP with SMA50 approaching SMA200.")

simple_table(
    ["Rule", "Requirement"],
    [
        ["SMA proximity", "SMA50 ≥ 92% of SMA200 (near_gc_pct: SMA50 within 5% of SMA200 from below)"],
        ["Price above SMAs", "Price above both SMA50 and SMA150"],
        ["RS condition", "Mansfield RS > 0"],
        ["Weekly RSI", "> 50"],
        ["VCP tightness", "ATR10 < ATR40 × 1.5 (vcp_atr_mult=1.5) — volatility contracting"],
        ["Pivot", "Price closes above 15-bar high (early_pivot_len=15)"],
        ["Volume", "≥1.5× average on breakout bar"],
        ["Stop", "Chandelier Exit: 3.5 × ATR14 below VCP low"],
        ["Targets", "T1 at 2.0R, T2 at 2.5R"],
    ],
    col_widths=[2.0, 5.0]
)
note("NOTE", "REV-EARLY signals are assigned priority 4 (highest) in the screener output. "
     "They represent the cleanest setups and should be taken first when capital is limited.")

divider()

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 7 — STOP-LOSS FRAMEWORK
# ═══════════════════════════════════════════════════════════════════════════
h1("Section 7 — Stop-Loss Framework")

h2("7.1 Chandelier Exit (Recovery)")
body("The primary stop mechanism for recovery trades uses a wider multiplier than bull market trades:")
p = doc.add_paragraph()
run = p.add_run("    Stop = Highest High (N bars) − (3.5 × ATR14)")
set_font(run, bold=True, size=11, color=C_AMBER)
body("The 3.5× multiplier (vs 3.0× for bull trades) reflects the higher volatility in recovery-phase stocks. "
     "A tighter stop during recovery will be stopped out by normal volatility before the trade develops.")

h2("7.2 Initial Stop Placement by Edge")
simple_table(
    ["Edge", "Initial Stop Rule"],
    [
        ["REV-CB", "Below the climax candle low − 0.5×ATR14"],
        ["REV-RS", "Chandelier Exit (3.5×ATR14) calculated from entry bar"],
        ["REV-EARLY", "Below VCP base low − 0.2×ATR14"],
    ],
    col_widths=[2.0, 5.0]
)

h2("7.3 Trailing the Stop (Weekly Review)")
body("After each weekly close, update the Chandelier Exit:")
bullet("Recalculate: Highest High (14 bars) − 3.5 × ATR14")
bullet("Move stop UP only — never move stop down")
bullet("Record the new SL in the Commander Web p_sl field (draggable price input)")
bullet("Dashboard TRADE STATUS will reflect updated SL in P&L display")

h2("7.4 Panic Exit Override")
body("If price closes below the 40-Week MA on a weekly candle, exit immediately regardless of "
     "Chandelier Exit level. Dashboard will display:")
p = doc.add_paragraph()
run = p.add_run('    "⚠️ STOP HIT (P&L%)"  or  "EXIT NOW (40WMA Break)"')
set_font(run, bold=True, size=11, color=C_RED)
note("CRITICAL", "40W MA break is a Weinstein sell signal for recovery trades. "
     "Do not argue with it. Exit at next market open.", color=C_RED)

h2("7.5 Noise Risk Check (ai_risk_manager.py)")
body("Before entering, verify the SL is not within normal noise range:")
bullet("get_noise_risk_stats() flags any position where SL < 1.5×ATR14")
bullet("If flagged, widen the stop or reduce size — never enter with an invisible stop")

divider()

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 8 — POSITION SIZING
# ═══════════════════════════════════════════════════════════════════════════
h1("Section 8 — Position Sizing")

h2("8.1 Risk-Per-Trade")
body("Maximum risk per recovery trade: 0.5%–1.0% of total trading capital. "
     "Use 0.5% for REV-CB (higher-risk bounce), 1.0% for REV-EARLY (highest quality).")

p = doc.add_paragraph()
run = p.add_run("    Shares = (Capital × Risk%) / (Entry Price − Stop Price)")
set_font(run, bold=True, size=11, color=C_NAVY)

h2("8.2 Adaptive ATR Multiplier (ai_risk_manager.py)")
body("get_adaptive_atr_multiplier() adjusts SL width by stock's Average Daily Range (ADR%):")
simple_table(
    ["ADR%", "ATR Multiplier", "Stock Type"],
    [
        ["< 1.5%", "1.5×", "Low-vol large cap — tight SL OK"],
        ["1.5%–3.0%", "2.0×", "Standard mid-cap"],
        ["> 3.0%", "2.5×", "High-vol small cap — needs more room"],
    ],
    col_widths=[1.2, 1.5, 3.5]
)

h2("8.3 Slot Limits")
simple_table(
    ["Condition", "Max New Slots"],
    [
        ["Market Health: BULLISH", "Up to 6 open positions"],
        ["Market Health: CORRECTION", "3–4 open positions max"],
        ["Market Health: RECOVERY/BEARISH", "2 open positions max, reduce sizing 50%"],
    ],
    col_widths=[3.0, 3.0]
)
note("RULE", "Never open a new recovery position when the Dashboard shows BEARISH market health.")

h2("8.4 Scaling In (REV-RS / REV-EARLY Only)")
body("If the Dashboard TRADE STATUS shows 'SCALE IN', criteria for adding to the position are met:")
bullet("Daily trend direction is up")
bullet("Daily RSI < 70 (not overbought)")
bullet("Price above 20 EMA")
bullet("No volatility squeeze active")
body("Add maximum 50% of original position size. Do not average up more than once.")

doc.add_page_break()

doc.save(OUTPUT)
print(f"Part 1 saved: {OUTPUT}")
print("Sections 1–8 complete.")
