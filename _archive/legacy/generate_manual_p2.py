"""
Recovery Trading Manual — Generator Part 2
Sections 9–15: Trade management, screening workflow, module ref sections
Appends to existing document.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = r"C:\Users\jayra\Documents\GeminiVSCode\Recovery_Trading_Manual.docx"

doc = Document(OUTPUT)

C_NAVY  = RGBColor(0x1F, 0x35, 0x64)
C_TEAL  = RGBColor(0x00, 0x70, 0x70)
C_AMBER = RGBColor(0xC5, 0x5A, 0x11)
C_GREEN = RGBColor(0x37, 0x86, 0x30)
C_LGREY = RGBColor(0xF2, 0xF2, 0xF2)
C_DGREY = RGBColor(0x40, 0x40, 0x40)
C_RED   = RGBColor(0xC0, 0x00, 0x00)

def set_font(run, bold=False, italic=False, size=11, color=C_DGREY):
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = color

def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    set_font(run, bold=True, size=18, color=C_NAVY)

def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    set_font(run, bold=True, size=14, color=C_TEAL)

def h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    set_font(run, bold=True, size=12, color=C_NAVY)

def body(text, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(indent * 0.25)
    run = p.add_run(text)
    set_font(run, size=10.5)

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_font(run, size=10.5)

def note(label, text, color=C_AMBER):
    p = doc.add_paragraph()
    r1 = p.add_run(f"{label}: ")
    set_font(r1, bold=True, size=10.5, color=color)
    r2 = p.add_run(text)
    set_font(r2, size=10.5)

def code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    set_font(run, bold=True, size=10, color=C_NAVY)

def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run("─" * 85)
    set_font(run, size=8, color=RGBColor(0xCC, 0xCC, 0xCC))

def shade_cell(cell, rgb):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:fill'), f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}")
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)

def simple_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        shade_cell(cell, C_NAVY)
        p = cell.paragraphs[0]
        run = p.add_run(h)
        set_font(run, bold=True, size=10, color=RGBColor(0xFF,0xFF,0xFF))
    for ri, row_data in enumerate(rows):
        row = t.add_row()
        bg = C_LGREY if ri % 2 == 0 else RGBColor(0xFF,0xFF,0xFF)
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            shade_cell(cell, bg)
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            set_font(run, size=10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 9 — TRADE MANAGEMENT
# ═══════════════════════════════════════════════════════════════════════════
h1("Section 9 — Trade Management")

h2("9.1 Dynamic Risk:Reward Targets")
simple_table(
    ["Market Phase", "T1 Target", "T2 Target", "Notes"],
    [
        ["Bull Market", "2.5R", "4.0R", "Can run full positions"],
        ["Recovery / Bear", "2.0R", "2.5R", "Take profits earlier — less runway"],
        ["REV-CB only", "2.0R", "N/A", "Single target — exit fully at T1"],
    ],
    col_widths=[2.0, 1.0, 1.0, 3.0]
)

h2("9.2 Partial Profit Protocol")
body("At T1: Sell 50% of position, move stop to breakeven on remainder.")
body("At T2: Sell remaining 50%. Close trade fully.")
note("RULE", "Never let a recovery trade turn from profit to loss. Once T1 is hit, "
     "stop MUST move to breakeven.", color=C_AMBER)

h2("9.3 Dashboard TRADE STATUS Logic")
body("The Dashboard shows one of four trade states for open positions:")
simple_table(
    ["Status", "Condition", "Action"],
    [
        ["⚠️ STOP HIT (P&L%)", "close < effectiveSL", "Exit immediately at next open"],
        ["TRIM (P&L%)", "RSI>75 OR dist>15% from EMA OR earnings <3 days", "Sell 25–50%, keep core"],
        ["SCALE IN (P&L%)", "Trend up + RSI<70 + above EMA20 + no squeeze", "Add up to 50% of original size"],
        ["HOLD (P&L%)", "All other in-profit or minor-loss states", "No action needed"],
    ],
    col_widths=[2.0, 3.0, 2.0]
)

h2("9.4 Time Stop")
body("Recovery trades have a maximum hold duration displayed as 'Time Stop' in the Dashboard. "
     "If the stock has not reached T1 within the time stop window, exit regardless of P&L. "
     "This prevents capital being tied up in stalled recovery setups.")
note("NOTE", "The Time Stop field appears in the Dashboard only when a position is open (p_entry > 0).")

h2("9.5 Earnings Risk Management")
body("When NEXT EARNINGS shows < 3 trading days:")
bullet("New entries: do not open unless you intend to hold through earnings")
bullet("Existing positions: Dashboard triggers TRIM signal automatically")
bullet("Consider closing fully if position is profitable and earnings are binary risk")

divider()

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 10 — PORTFOLIO MANAGEMENT
# ═══════════════════════════════════════════════════════════════════════════
h1("Section 10 — Portfolio Management")

h2("10.1 Portfolio Slots")
body("Recovery trades are limited to 2–6 concurrent open positions depending on market health. "
     "Each slot is tracked in the Commander Web interface (weinstein_commander_web_v3.0.py).")

simple_table(
    ["Market Health", "Max Slots", "Max Risk per Trade"],
    [
        ["BULLISH (Confirmed)", "6", "1.0% of capital"],
        ["CORRECTION (In Uptrend)", "4", "0.75% of capital"],
        ["RECOVERY (Weak)", "2", "0.5% of capital"],
        ["BEARISH", "0 new entries", "Exit only"],
    ],
    col_widths=[2.5, 1.5, 2.5]
)

h2("10.2 Portfolio Vitals (ai_risk_manager.py)")
body("calculate_portfolio_vitals() returns the following metrics from closed trade history:")
simple_table(
    ["Metric", "Description", "Target"],
    [
        ["Sharpe Ratio", "Risk-adjusted return", "> 1.5"],
        ["Max Drawdown", "Peak-to-trough equity loss", "< 15%"],
        ["Calmar Ratio", "Return / Max Drawdown", "> 1.0"],
        ["Expectancy", "Avg profit per trade in R", "> 0.3R"],
        ["Win Rate", "% of winning trades", "> 45%"],
        ["Avg Win / Avg Loss", "Reward vs loss ratio", "> 2.0"],
        ["Total Return %", "Cumulative return on capital", "Benchmark: CNX500"],
        ["Unrealized P&L", "Mark-to-market open positions", "Monitor weekly"],
    ],
    col_widths=[1.8, 3.0, 1.8]
)

h2("10.3 Correlation Risk")
body("get_portfolio_correlation_matrix() identifies 'shadow pairs' — positions with correlation > 0.85.")
bullet("Shadow pairs move together: if one stops out, the other likely will too")
bullet("Do not hold more than one shadow pair simultaneously")
bullet("Diversity score is returned — aim for score > 0.6")

h2("10.4 Weekly SL Update Workflow")
body("Every Sunday after market close:")
bullet("Step 1: Open Commander Web → Portfolio tab")
bullet("Step 2: For each open position, recalculate Chandelier Exit: Highest High (14d) − 3.5 × ATR14")
bullet("Step 3: If new SL > current SL → update p_sl in Dashboard (draggable price handle on chart)")
bullet("Step 4: Verify Dashboard shows updated TRADE STATUS with new P&L reference")
bullet("Step 5: Log the change in the trade journal")
note("RULE", "SL moves UP only. Widening a stop to 'give the trade more room' is never acceptable.")

h2("10.5 P&L Tracking")
body("The Dashboard shows real-time P&L% whenever p_entry and p_sl are set:")
code("P&L% = ((current_close - entry_price) / entry_price) × 100")
body("The Trade Status label appends this value: e.g., 'HOLD (+7.3%)' or 'TRIM (+18.2%)'.")
body("For closed trade performance tracking, export from Commander Web to the closed_trades CSV "
     "and run calculate_portfolio_vitals() for full statistics.")

divider()

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 11 — SCREENING WORKFLOW
# ═══════════════════════════════════════════════════════════════════════════
h1("Section 11 — Daily/Weekly Screening Workflow")

h2("11.1 Pre-Market Regime Check (Daily, 5 min)")
bullet("Step 1: Check CNX500 status in Dashboard (Market Health row)")
bullet("Step 2: If BEARISH → no new entries today. Review exits only.")
bullet("Step 3: If CORRECTION or RECOVERY → reduce size, tighten criteria")
bullet("Step 4: Note if any open positions show ⚠️ STOP HIT from overnight move")

h2("11.2 Chartink Scanner Run (Weekly, Sunday)")
body("Run chartink_scanner_pro.py to pull fresh candidates from three recovery scanners:")
simple_table(
    ["Scanner", "Output File", "Edge"],
    [
        ["Scanner 5 — REV-RS", "Recovery_RS_Survivors.csv", "RS positive + price recovering"],
        ["Scanner 6 — REV-CB", "Recovery_Climax_Bounce.csv", "Stretched below SMA200, RS positive"],
        ["Scanner 7 — REV-EARLY", "Recovery_Early_Birds.csv", "Near golden cross, VCP tightening"],
    ],
    col_widths=[2.5, 2.5, 2.5]
)
note("NOTE", "Chartink scanners apply broad filters. The recovery_screener.py step below applies "
     "the full RFF + regime + composite score.")

h2("11.3 Recovery Screener Deep Score (Weekly)")
body("Run recovery_screener.py on the Chartink output:")
bullet("Input: the three CSV files from Chartink")
bullet("Output: Recovery_Screener_Results.csv sorted by Signal (desc) then Score (desc)")
bullet("Signal values: 4=REV-EARLY, 3=REV-RS, 2=REV-CB, 1=CB-Watch, 0=None")
bullet("Signal hold window: 5 trading days — signals remain visible for weekend review")
note("CONFIG", "Adjust rff_min_score, near_gc_pct, and vcp_atr_mult in the CONFIG block "
     "at the top of recovery_screener.py to tune strictness.")

h2("11.4 AI Grading (On Shortlist)")
body("For the top 10–15 candidates from the screener, run get_weinstein_score() "
     "from ai_grading_engine.py:")
bullet("Grading is cached — first run fetches data, subsequent runs use cache")
bullet("Focus on A and B grade stocks (score ≥ 60)")
bullet("Check the catalyst tag: POS-ACCUM, POS-BO, SWG-BO are highest priority")
bullet("SWG-PB (pullback) is valid for REV-EARLY; SWG-REV for REV-CB")

h2("11.5 Dashboard Review (On Watchlist)")
body("Load each shortlisted symbol in TradingView with Dashboard v63.3 active:")
bullet("Check: Stage Display shows Stage 1 (Constructive) or early Stage 2")
bullet("Check: Action Signal ≥ 3 stars out of 7")
bullet("Check: Alpha Score ≥ 50")
bullet("Check: Room for Trade — T1 not blocked by overhead resistance")
bullet("Check: Sector Velocity — prefer ACCELERATING LEADER or HIDDEN ACCUMULATION")
bullet("Check: Next Earnings — not within 3 days")

h2("11.6 Entry Decision Gate (Checklist)")
simple_table(
    ["Gate", "Check", "Pass Condition"],
    [
        ["1", "Market Health", "Not BEARISH"],
        ["2", "Regime Gate", "Three-Way Gate not fully blocked"],
        ["3", "RFF", "≥ rff_min_score checks passed"],
        ["4", "Composite Score", "≥ 8 (medium conviction minimum)"],
        ["5", "AI Grade", "B or A (score ≥ 60)"],
        ["6", "Stage", "Stage 1 Constructive or Stage 2"],
        ["7", "Action Signal", "≥ 3 stars"],
        ["8", "Room for Trade", "T1 not blocked"],
        ["9", "Earnings", "Not within 3 days"],
        ["10", "Slot available", "Portfolio not at max slots"],
    ],
    col_widths=[0.5, 2.5, 3.5]
)
note("RULE", "All 10 gates must pass before entry. A 9/10 is still a no.")

divider()

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 12 — ALGO SIGNAL MODE
# ═══════════════════════════════════════════════════════════════════════════
h1("Section 12 — Algo Signal Mode (Dashboard v63.3)")

h2("12.1 What It Does")
body("When a catalyst fires on a non-slot stock (no p_entry set), the Dashboard auto-draws "
     "SL, T1, and T2 levels on the chart using Pine Script line.new() calls. "
     "This gives an instant visual risk/reward map without manually entering trade parameters.")

h2("12.2 SL and Target Formulas")
simple_table(
    ["Level", "Formula", "Notes"],
    [
        ["SL (Positional)", "d_l20 − (d_atr10 × 0.2)", "Low of last 20 daily bars minus 20% of ATR10"],
        ["SL (Swing)", "d_l20 × 0.998", "0.2% below 20-bar low"],
        ["T1", "entry + risk × 2.0", "2R target"],
        ["T2", "entry + risk × 3.5", "3.5R target"],
    ],
    col_widths=[1.8, 3.0, 2.2]
)
note("NOTE", "These lines are drawn automatically when a catalyst fires. "
     "They are visual guides — use them as a reference for manual entry decisions.")

h2("12.3 Catalyst Tags and Priority")
simple_table(
    ["Tag", "Meaning", "Dashboard Shows"],
    [
        ["POS-ACCUM", "Stage 1/2, OBV breaking 50-bar high, VCP tight, flat base", "⚡ Positive Accumulation"],
        ["POS-BO", "Stage 2, near 52W high, institutional volume surge", "🚀 Positional Breakout"],
        ["SWG-GAP", "Stage 2, gap-up open above prior high close", "🔼 Gap & Go"],
        ["SWG-BO", "Stage 2, VCP tight, closes above 10-bar high, above-avg vol", "📈 Swing Breakout"],
        ["SWG-PB", "Stage 2, close to EMA20, low vol pullback, RSI healthy", "↩️ Swing Pullback"],
        ["SWG-REV", "Stage 2, RSI3 < 20, stretched below EMA20", "🔄 Mean Reversion"],
        ["NONE", "No catalyst active", "—"],
    ],
    col_widths=[1.5, 3.0, 2.5]
)

divider()

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 13 — READING THE DASHBOARD
# ═══════════════════════════════════════════════════════════════════════════
h1("Section 13 — Reading the Dashboard at a Glance")

h2("13.1 Alpha Score Interpretation")
body("The Alpha Score (0–100) summarises overall setup quality in four weighted components:")
simple_table(
    ["Component", "Weight", "What It Measures"],
    [
        ["Trend Structure", "30 pts", "SMA alignment, stage, 30W MA slope"],
        ["Momentum", "30 pts", "RSI, EMA proximity, price action structure"],
        ["Volume", "20 pts", "Relative volume, OBV, accumulation patterns"],
        ["Safety", "20 pts", "Volatility state, SL distance, liquidity"],
    ],
    col_widths=[2.0, 1.2, 3.8]
)

simple_table(
    ["Score", "Grade", "Meaning"],
    [
        ["80–100", "A+", "Institutional quality — highest priority"],
        ["60–79", "B", "Good setup — standard entry"],
        ["< 60", "C", "Marginal — skip or reduce size significantly"],
    ],
    col_widths=[1.2, 0.8, 4.5]
)

h2("13.2 Action Signal (7-Component Star Rating)")
body("The Action Signal rates the current setup across 7 binary components. "
     "Each component that passes adds one star (max 7★):")
simple_table(
    ["Component", "Pass Condition"],
    [
        ["Trend", "Daily and weekly trend both up"],
        ["Momentum", "RSI in healthy range (not overbought/oversold)"],
        ["Volume", "Above-average volume on recent up-days"],
        ["VP (POC)", "Price above Volume Profile Point of Control"],
        ["Price Action", "Higher highs / higher lows pattern intact"],
        ["ADX", "ADX > 20 (trend has conviction, not choppy)"],
        ["Liquidity Sweep", "Recent stop-hunt below support followed by recovery"],
    ],
    col_widths=[2.0, 5.0]
)
note("GUIDE", "For recovery trades: aim for ≥ 3 stars. 5+ stars is high-conviction entry.")

h2("13.3 Volatility State")
simple_table(
    ["State", "Condition", "Trading Implication"],
    [
        ["VOLATILITY SQUEEZE (Ready)", "ATR5 < ATR20 × 0.60", "Breakout imminent — ideal entry zone"],
        ["VCP (True Dry Up)", "ATR10 < ATR40 × 0.75 AND relVol < 0.8", "Classic VCP — highest quality base"],
        ["VCP (Actionable)", "ATR10 tight but vol not yet dry", "Good base — watch for volume expansion"],
        ["Vol Expansion", "ATR expanding vs recent average", "Breakout in progress or failed — confirm direction"],
        ["Normal Volatility", "None of the above", "Standard conditions"],
    ],
    col_widths=[2.2, 2.5, 2.5]
)

h2("13.4 Volume State")
simple_table(
    ["State", "Condition", "Meaning"],
    [
        ["ACCUMULATION (x.x×)", "relVol > 1.5 + up close", "Institutions buying — bullish"],
        ["DISTRIBUTION (x.x×)", "relVol > 1.5 + down close", "Institutions selling — bearish"],
        ["DRY UP (x.x×)", "relVol < 0.7", "No selling pressure — base forming"],
        ["NORMAL (x.x×)", "All other", "No extreme reading"],
    ],
    col_widths=[2.2, 2.5, 2.5]
)

h2("13.5 EMA Proximity State")
simple_table(
    ["State", "Meaning"],
    [
        ["RESISTANCE (Hit Head)", "Price touched EMA20 and was rejected — bearish near-term"],
        ["SUPPORT (Bounced)", "Price touched EMA20 and bounced — bullish entry signal"],
        ["NEAR (x.x%)", "Price within x% of EMA20 — monitor for bounce or break"],
        ["PRICE ABOVE", "Price comfortably above EMA20 — trend intact"],
        ["PRICE BELOW", "Price below EMA20 — caution, trend weakening"],
    ],
    col_widths=[2.5, 4.5]
)

divider()

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 14 — SECTOR ANALYSIS
# ═══════════════════════════════════════════════════════════════════════════
h1("Section 14 — Sector Analysis")

h2("14.1 Sector Velocity")
body("Sector Velocity measures the rate-of-change of the sector RS slope — "
     "how fast the sector is gaining or losing relative strength vs the market.")
simple_table(
    ["State", "Meaning", "Trade Bias"],
    [
        ["🔥 ACCELERATING LEADER", "Sector RS slope rising fast", "Highest priority — tide is with you"],
        ["🛑 EXHAUSTED LEADER", "Sector was leading but slope flattening", "Reduce new entries, protect profits"],
        ["⚡ HIDDEN ACCUMULATION", "Sector weak overall but RS slope turning up", "Early positioning opportunity"],
        ["🧊 DEAD MONEY", "Sector lagging and RS slope falling", "Avoid new entries"],
        ["🧊 Static", "No meaningful change in RS slope", "Neutral — stock-specific moves only"],
    ],
    col_widths=[2.5, 2.5, 2.0]
)

h2("14.2 Multi-RS View")
body("The Dashboard shows three RS readings:")
bullet("RS (vs Nifty 50) — performance vs the large-cap index")
bullet("RS (vs N500) — performance vs the broad market (primary Mansfield RS)")
bullet("RS (sector) — how the stock ranks within its own sector")
body("A stock with positive RS vs N500 but negative RS vs its sector may indicate "
     "broad market exposure rather than sector-specific strength.")

h2("14.3 Sector Stage (Weekly)")
body("The sector itself is classified using the same Weinstein stage analysis as individual stocks. "
     "Prefer stocks in sectors showing Stage 2 or early Stage 1 Constructive on the weekly chart.")
note("RULE", "Never take a recovery trade in a Stage 4 sector. The sector headwind will "
     "overwhelm individual stock technical setups.")

divider()

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 15 — JOURNAL AND REVIEW PROTOCOL
# ═══════════════════════════════════════════════════════════════════════════
h1("Section 15 — Journal and Review Protocol")

h2("15.1 Trade Entry Journal Template")
simple_table(
    ["Field", "Example"],
    [
        ["Date", "2026-04-14"],
        ["Symbol", "RELIANCE"],
        ["Edge", "REV-EARLY"],
        ["Composite Score", "14/20"],
        ["AI Grade", "B (68/100)"],
        ["Alpha Score", "72"],
        ["Action Signal", "5/7 ★"],
        ["Entry Price", "₹1,285"],
        ["Stop Loss", "₹1,240"],
        ["T1", "₹1,375 (2.0R)"],
        ["T2", "₹1,420 (2.5R)"],
        ["Position Size", "78 shares (₹1L capital, 0.75% risk)"],
        ["ATR14", "₹22.4"],
        ["Chandelier SL", "₹1,232 (check vs initial SL)"],
        ["Regime", "CNX500 CORRECTION — 2 slots max"],
        ["Sector", "Energy — ACCELERATING LEADER"],
        ["Thesis", "Near golden cross, VCP 18 days, RS positive, earnings 45 days away"],
    ],
    col_widths=[2.5, 4.5]
)

h2("15.2 Weekly Review Questions")
bullet("Did any stops trigger? Were they placed correctly per Chandelier rules?")
bullet("Are any positions showing TRIM signal? Have I acted?")
bullet("Has market health changed? Do I need to reduce slots?")
bullet("Are any Time Stops approaching? Do I exit or extend?")
bullet("What is the portfolio diversity score? Any shadow pairs?")
bullet("Did the screener produce new high-score candidates? Do I have open slots?")

h2("15.3 Monthly Performance Review")
body("Run calculate_portfolio_vitals() with the closed trade CSV. Review:")
bullet("Sharpe ratio trend — improving or degrading?")
bullet("Win rate by edge type — which edge is performing best this quarter?")
bullet("Max drawdown — are position sizes appropriate for current volatility?")
bullet("Expectancy — is the system profitable on a per-trade basis?")
body("If expectancy drops below 0.2R over 20+ trades, pause and review entry criteria.")

doc.add_page_break()
doc.save(OUTPUT)
print(f"Part 2 saved: {OUTPUT}")
print("Sections 9–15 complete.")
